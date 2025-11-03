import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

doc = Document(r"C:\Users\Marco.Africani\Desktop\Month recap\month recap_EUR.docx")

print("="*80)
print("VERIFYING SPECIFIC FIXES")
print("="*80)

# Test 1: Lafleur 2022
print("\n1. LAFLEUR 2022:")
for i, para in enumerate(doc.paragraphs):
    if 'Lafleur' in para.text and '2022' in para.text:
        print(f"[Para {i}] {para.text[:200]}")
        print("Expected: 1150 CHF → 1250 EUR (rounded from 1242 to 1250)")
        for run in para.runs:
            if run.font.highlight_color:
                print(f"  Highlighted: '{run.text}' [{run.font.highlight_color}]")
        break

# Test 2: Harlan Estate 2021
print("\n2. HARLAN ESTATE 2021:")
for i, para in enumerate(doc.paragraphs):
    if 'Harlan Estate' in para.text and '2021' in para.text:
        print(f"[Para {i}] {para.text[:200]}")
        print("Expected: Should use last row (row 809) = 1300 CHF → 1405 EUR (rounded to 1405)")
        for run in para.runs:
            if run.font.highlight_color:
                print(f"  Highlighted: '{run.text}' [{run.font.highlight_color}]")
        break

# Test 3: L'Évangile 2015
print("\n3. L'ÉVANGILE 2015:")
for i, para in enumerate(doc.paragraphs):
    if 'vangile' in para.text.lower() and '2015' in para.text:
        print(f"[Para {i}] {para.text[:300]}")
        print("Expected:")
        print("  - 195 CHF (normal) → 210 EUR")
        print("  - 190 CHF (36x) → 205 EUR")
        for run in para.runs:
            if run.font.highlight_color:
                print(f"  Highlighted: '{run.text}' [{run.font.highlight_color}]")
        break

# Test 4: Rieussec (previous fix verification)
print("\n4. RIEUSSEC 2019 (Previous fix):")
for i, para in enumerate(doc.paragraphs):
    if 'Rieussec' in para.text and 'sauternes' in para.text.lower():
        print(f"[Para {i}] {para.text[:300]}")
        print("Expected:")
        print("  - 100 CHF (market) → 108 EUR [RED]")
        print("  - 42 CHF (normal) → 45 EUR")
        print("  - 39 CHF (36 bottles) → 43 EUR")
        break

print("\n" + "="*80)
