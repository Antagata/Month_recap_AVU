#!/usr/bin/env python3
"""
DeepL Text Translation Script for Word Documents
Extracts text from Word, translates via DeepL API, creates new Word documents
"""

import os
import sys
import io
import time
import requests
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX

# Fix encoding for Windows console
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

# Configuration
DEEPL_API_KEY = "374a8965-101a-4538-bc65-54506552650e"
SOURCE_FILE = r"C:\Users\Marco.Africani\Desktop\Month recap\month recap_EUR.docx"
OUTPUT_DIR = r"C:\Users\Marco.Africani\Desktop\Month recap\Translations"

# Target languages
LANGUAGES = {
    'DE': 'German',
    'ES': 'Spanish',
    'FR': 'French',
    'IT': 'Italian',
    'PT-PT': 'Portuguese',
    'RU': 'Russian'
}

# DeepL API endpoint for text translation
DEEPL_API_URL = "https://api-free.deepl.com/v2/translate"


def translate_text(text, target_lang):
    """Translate text using DeepL API"""
    if not text or not text.strip():
        return text

    data = {
        'auth_key': DEEPL_API_KEY,
        'text': text,
        'target_lang': target_lang,
        'source_lang': 'EN',
        'preserve_formatting': '1'
    }

    response = requests.post(DEEPL_API_URL, data=data)

    if response.status_code == 200:
        result = response.json()
        return result['translations'][0]['text']
    else:
        print(f"      Warning: Translation failed for text snippet: {response.status_code}")
        return text  # Return original if translation fails


def translate_document(source_path, target_lang, output_path):
    """
    Translate a Word document by extracting text, translating, and rebuilding
    """
    print(f"\n{'='*80}")
    print(f"Translating to {LANGUAGES[target_lang]} ({target_lang})")
    print(f"{'='*80}")

    try:
        # Load source document
        print(f"   📄 Loading source document...")
        doc = Document(source_path)

        # Create new document for translation
        translated_doc = Document()

        # Note: Styles cannot be copied directly in python-docx
        # The new document will use default styles

        # Translate each paragraph
        total_paragraphs = len(doc.paragraphs)
        print(f"   🔄 Translating {total_paragraphs} paragraphs...")

        for i, paragraph in enumerate(doc.paragraphs):
            if i % 10 == 0 and i > 0:
                print(f"      Progress: {i}/{total_paragraphs} paragraphs...")

            # Create new paragraph
            new_para = translated_doc.add_paragraph()

            # Copy paragraph style
            if paragraph.style:
                try:
                    new_para.style = paragraph.style
                except:
                    pass  # Skip if style doesn't exist

            # Translate and add runs
            if paragraph.text.strip():
                # Translate entire paragraph text
                translated_text = translate_text(paragraph.text, target_lang)

                # Add as single run (preserving basic formatting)
                new_run = new_para.add_run(translated_text)

                # Try to copy formatting from first run of original
                if paragraph.runs and len(paragraph.runs) > 0:
                    orig_run = paragraph.runs[0]
                    if orig_run.bold:
                        new_run.bold = True
                    if orig_run.italic:
                        new_run.italic = True
                    if orig_run.underline:
                        new_run.underline = True
                    if orig_run.font.size:
                        new_run.font.size = orig_run.font.size
                    if orig_run.font.color and orig_run.font.color.rgb:
                        new_run.font.color.rgb = orig_run.font.color.rgb

                # Small delay to avoid rate limiting
                time.sleep(0.1)

        # Handle tables (if any)
        if doc.tables:
            print(f"   📊 Translating {len(doc.tables)} tables...")
            for table_idx, table in enumerate(doc.tables):
                new_table = translated_doc.add_table(rows=len(table.rows), cols=len(table.columns))

                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        cell_text = cell.text
                        if cell_text.strip():
                            translated_cell_text = translate_text(cell_text, target_lang)
                            new_table.rows[row_idx].cells[col_idx].text = translated_cell_text
                            time.sleep(0.1)

        # Save translated document
        print(f"   💾 Saving translated document...")
        translated_doc.save(output_path)
        print(f"   ✅ Saved to: {output_path}")

        return True

    except Exception as e:
        print(f"   ❌ Error: {e}")
        import traceback
        traceback.print_exc()
        return False


def main():
    """Main translation function"""
    print("\n" + "="*80)
    print("DeepL Text Translation for Word Documents")
    print("="*80)

    # Check if source file exists
    if not os.path.exists(SOURCE_FILE):
        print(f"\n❌ Error: Source file not found: {SOURCE_FILE}")
        print("   Please make sure the converted EUR document exists.")
        return

    # Create output directory if it doesn't exist
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    print(f"\n✅ Source file: {SOURCE_FILE}")
    print(f"✅ Output directory: {OUTPUT_DIR}")
    print(f"✅ Languages to translate: {len(LANGUAGES)}")

    # Translate to each language
    successful = 0
    failed = 0

    for lang_code, lang_name in LANGUAGES.items():
        # Generate output filename
        output_filename = f"month recap_{lang_code}.docx"
        output_path = os.path.join(OUTPUT_DIR, output_filename)

        # Translate
        if translate_document(SOURCE_FILE, lang_code, output_path):
            successful += 1
        else:
            failed += 1

        # Delay between documents to avoid rate limiting
        if lang_code != list(LANGUAGES.keys())[-1]:  # Not the last one
            print(f"\n   ⏳ Waiting 3 seconds before next translation...")
            time.sleep(3)

    # Summary
    print("\n" + "="*80)
    print("TRANSLATION SUMMARY")
    print("="*80)
    print(f"✅ Successful translations: {successful}/{len(LANGUAGES)}")
    if failed > 0:
        print(f"❌ Failed translations: {failed}/{len(LANGUAGES)}")

    print(f"\nAll translated documents saved to:")
    print(f"  {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
