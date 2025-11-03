# ==============================================================================
# SCRIPT NAME: word_converter.py
# DESCRIPTION: Converts CHF numbers to EUR in a Word document based on an Excel map.
#              Applies 1.08 fallback with integer floor rounding if a number is
#              not found. Highlights converted numbers based on conversion type
#              (Red for 1.08 fallback, Yellow for duplicates).
# ==============================================================================

import pandas as pd
from docx import Document
import re
from docx.enum.text import WD_COLOR_INDEX
import math
from collections import Counter

# --- CONFIGURATION (UPDATE THESE PATHS) ---
# NOTE: Use raw strings (r"...") for Windows paths
WORD_FILE_PATH = r"C:\Users\Marco.Africani\Desktop\Month recap\month recap.docx"
EXCEL_FILE_PATH = r"C:\Users\Marco.Africani\Desktop\Month recap\Conversion_month.xlsx"
NEW_WORD_FILE_PATH = r"C:\Users\Marco.Africani\Desktop\Month recap\month recap_EUR.docx"

# Define the correct column names from your Excel file's header
CHF_COL = 'Unit Price'        # Corresponds to Excel Column I
EUR_COL = 'Unit Price (EUR)'  # Corresponds to Excel Column J

# Pre-defined list of duplicate CHF values from the previous run
# The converted EUR value corresponding to these originals will be highlighted YELLOW.
DUPLICATE_CHFS = {
    '39.00', '42.00', '50.00', '165.00', '195.00', 
    '33.00', '34.00', '79.00', '36.00', '150.00'
}

# Define Regex Patterns
NUMBER_PATTERN = r'\d+\.\d{2}' # Matches numbers with two decimal places (e.g., '25.00')
CHF_PATTERN = r'[Cc][Hh][Ff]' # Matches CHF, Chf, chf, etc.


def load_data_and_document():
    """Loads the Excel data, creates the conversion map, and loads the Word document."""
    conversion_map = {}
    doc = None
    
    # 1. Load Excel File and Create Conversion Dictionary
    try:
        df = pd.read_excel(EXCEL_FILE_PATH)
        
        # Standardize CHF column to be keys formatted as 'X.XX'
        df['CHF_KEY_FORMATTED'] = df[CHF_COL].astype(float).round(2).apply(lambda x: f'{x:.2f}')
        
        # Standardize EUR column to be values formatted as 'X.XX'
        df['EUR_VALUE_FORMATTED'] = df[EUR_COL].astype(float).round(2).apply(lambda x: f'{x:.2f}')
        
        conversion_map = pd.Series(df['EUR_VALUE_FORMATTED'].values, index=df['CHF_KEY_FORMATTED']).to_dict()
        
        print("✅ Excel file loaded and conversion map created.")
        
    except FileNotFoundError:
        print(f"❌ Error: Excel file not found at {EXCEL_FILE_PATH}. Check path/file name.")
    except KeyError as e:
        print(f"❌ Error: Column not found. Check if column headers are '{CHF_COL}' and '{EUR_COL}'. Details: {e}")
    except Exception as e:
        print(f"❌ An unexpected error occurred while loading the Excel file: {e}")

    # 2. Load Word Document
    try:
        doc = Document(WORD_FILE_PATH)
        print(f"✅ Word file loaded from {WORD_FILE_PATH}")
    except FileNotFoundError:
        print(f"❌ Error: Word file not found at {WORD_FILE_PATH}. Check path/file name.")
    except Exception as e:
        print(f"❌ An unexpected error occurred while loading the Word file: {e}")

    return doc, conversion_map


def replace_and_highlight(paragraph, conversion_map, duplicates, all_numbers_found):
    """
    Performs search and replace at the run level, applying red (1.08 fallback) 
    or yellow (duplicate) highlighting.
    """
    
    local_replacements = 0 
    
    # 1. Replace CHF with EUR in the paragraph text
    paragraph.text = re.sub(CHF_PATTERN, 'EUR', paragraph.text)
    
    text = paragraph.text
    matches = re.findall(NUMBER_PATTERN, text)
    
    # Dictionary of replacements: {chf_str: (new_eur_str, highlight_color_index)}
    replacements_to_do = {}
    for chf_str in set(matches):
        
        # --- VINTAGE/YEAR SKIP CHECK (e.g., skips 2019.00) ---
        try:
            float_val = float(chf_str)
            if float_val.is_integer() and 1000 <= int(float_val) <= 9999:
                continue
        except ValueError:
            continue
        # ----------------------------------------------------
        
        all_numbers_found.append(chf_str) 

        eur_value = None
        highlight_color = None
        
        if chf_str in conversion_map:
            # Case 1: Found in Excel map (no red highlight)
            eur_value = conversion_map[chf_str]
        else:
            # Case 2: Not found in map, apply 1.08 multiplier (RED highlight required)
            try:
                raw_eur = float(chf_str) * 1.08
                # Apply floor rounding and format with .00
                floored_eur = math.floor(raw_eur)
                eur_value = f"{floored_eur:.2f}"
                
                highlight_color = WD_COLOR_INDEX.RED # Set RED highlight flag
            except ValueError:
                continue
        
        # Final highlight decision
        if eur_value:
            # If it's a duplicate AND not already marked RED, mark it YELLOW
            if chf_str in duplicates and highlight_color != WD_COLOR_INDEX.RED:
                highlight_color = WD_COLOR_INDEX.YELLOW
                
            replacements_to_do[chf_str] = (eur_value, highlight_color)


    # 2. Perform number replacement with highlighting (rebuilding runs)
    new_text = text
    
    # Generate the final text string with all replacements
    for chf_str, (eur_value, _) in replacements_to_do.items():
        new_text = new_text.replace(chf_str, eur_value)
        
    # Delete existing runs
    for i in range(len(paragraph.runs) - 1, -1, -1):
        p = paragraph._element
        p.remove(paragraph.runs[i]._element)

    # Rebuild the paragraph content, segment by segment
    current_position = 0
    eur_to_chf_map = {v[0]: (k, v[1]) for k, v in replacements_to_do.items()}
    
    # Find all occurrences of the new EUR values in the new_text
    positions = []
    for eur_value in eur_to_chf_map.keys():
        start = 0
        while True:
            start = new_text.find(eur_value, start)
            if start == -1:
                break
            positions.append((start, eur_value))
            start += len(eur_value)
    
    positions.sort(key=lambda x: x[0])
    
    
    for start_index, eur_value in positions:
        
        # 1. Add text segment BEFORE the replaced value
        text_before = new_text[current_position:start_index]
        if text_before:
            paragraph.add_run(text_before)
        
        # 2. Add the HIGHLIGHTED/non-highlighted new value
        _, highlight_color = eur_to_chf_map[eur_value]
        
        run = paragraph.add_run(eur_value)
        
        if highlight_color is not None:
            run.font.highlight_color = highlight_color
        
        local_replacements += 1
        current_position = start_index + len(eur_value)
    
    # 3. Add any remaining text after the last replacement
    if current_position < len(new_text):
        paragraph.add_run(new_text[current_position:])
        
    return local_replacements


def main():
    """Main function to orchestrate the conversion process."""
    
    # Ensure necessary packages are installed
    # In VS Code terminal: pip install pandas python-docx
    
    doc, conversion_map = load_data_and_document()

    if not doc or not conversion_map:
        print("\nOperation aborted due to file loading errors.")
        return

    total_replacements = 0
    all_numbers_found = []

    # 1. Iterate through all paragraphs in the document
    for paragraph in doc.paragraphs:
        total_replacements += replace_and_highlight(paragraph, conversion_map, DUPLICATE_CHFS, all_numbers_found)

    # 2. Iterate through all tables in the document (if any)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    total_replacements += replace_and_highlight(paragraph, conversion_map, DUPLICATE_CHFS, all_numbers_found)


    # --- Duplicate Check Report ---
    duplicate_counts = Counter(all_numbers_found)
    duplicates_found_in_doc = {num: count for num, count in duplicate_counts.items() if count > 1}
    
    print("\n--- Duplicate Number Check (Original CHF Values) ---")
    if duplicates_found_in_doc:
        print("⚠️ **Duplicate numbers found** in the original Word file:")
        for num, count in duplicates_found_in_doc.items():
            print(f"- {num} appears {count} times (highlighted in YELLOW if not converted by 1.08).")
    else:
        print("✅ No duplicate numbers found in the original Word file.")
    print("----------------------------------------------------")
    
    # 3. Save the new document
    try:
        doc.save(NEW_WORD_FILE_PATH)
        print(f"\n🎉 **Success!** New Word file saved as: **{NEW_WORD_FILE_PATH}**")
        print(f"Total numbers replaced and processed: **{total_replacements}**")
        print("NOTE: Numbers converted with the 1.08 factor are rounded down to the nearest integer and highlighted in RED.")
    except Exception as e:
        print(f"\n❌ Error saving the new Word file: {e}")


if __name__ == "__main__":
    main()

# ==============================================================================