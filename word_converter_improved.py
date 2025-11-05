# ==============================================================================
# SCRIPT NAME: word_converter_improved.py
# DESCRIPTION: Converts CHF numbers to EUR in a Word document based on an Excel map.
#              Uses fuzzy wine name matching to resolve duplicate CHF prices.
#              Applies 1.08 fallback with integer floor rounding if a number is
#              not found. Highlights converted numbers based on conversion type:
#              - RED: 1.08 fallback conversion
#              - YELLOW: Duplicate CHF with ambiguous EUR conversion (not in Excel)
#              - GREEN: Successfully matched using wine name proximity
# ==============================================================================

import pandas as pd
from docx import Document
import re
from docx.enum.text import WD_COLOR_INDEX
import math
from collections import Counter, defaultdict
from difflib import SequenceMatcher
import sys
import io

# Fix encoding for Windows console
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

# --- CONFIGURATION (UPDATE THESE PATHS) ---
WORD_FILE_PATH = r"C:\Users\Marco.Africani\Desktop\Month recap\month recap.docx"
EXCEL_FILE_PATH = r"C:\Users\Marco.Africani\Desktop\Month recap\Conversion_month.xlsx"
NEW_WORD_FILE_PATH = r"C:\Users\Marco.Africani\Desktop\Month recap\month recap_EUR.docx"
LINES_EXCEL_PATH = r"C:\Users\Marco.Africani\Desktop\Month recap\Lines.xlsx"

# Excel column names
CHF_COL = 'Unit Price'
EUR_COL = 'Unit Price (EUR)'
WINE_NAME_COL = 'Wine Name'
CAMPAIGN_SUBTYPE_COL = 'Campaign Sub-Type'
CAMPAIGN_TYPE_COL = 'Campaign Type'
SIZE_COL = 'Size'
MIN_QUANTITY_COL = 'Minimum Quantity'
COMPETITOR_CODE_COL = 'Competitor Code'
PRODUCER_NAME_COL = 'Producer Name'
VINTAGE_COL = 'Vintage'
ITEM_NO_COL = 'Item No.'

# Define Regex Patterns
# Match numbers with decimals, including Swiss format with apostrophe or curly quote (e.g., 1'500.00, 1'500.00)
NUMBER_PATTERN = r"\d+(?:['\u2019]\d{3})*\.\d{2}"
CHF_PATTERN = r'[Cc][Hh][Ff]'

# Additional pattern to catch numbers like "CHF 100" or "CHF 42+vat" (without decimals)
CHF_NUMBER_NO_DECIMAL = r"[Cc][Hh][Ff]\s+(\d+(?:['\u2019]\d{3})*)(?![.\d])"

# Pattern to catch "NUMBER CHF" format without decimals (e.g., "190 CHF")
NUMBER_THEN_CHF = r"(\d+(?:['\u2019]\d{3})*)\s+[Cc][Hh][Ff]"

# Pattern to catch "NUMBER.XXCHF" format (no space before CHF, e.g., "33.00CHF")
NUMBER_NOSPACE_CHF = r"(\d+(?:['\u2019]\d{3})*\.\d{2})[Cc][Hh][Ff]"

# Fuzzy matching threshold (0-1, where 1 is exact match)
FUZZY_MATCH_THRESHOLD = 0.5


def normalize_wine_name(name):
    """
    Normalize wine name for better matching:
    - Convert to lowercase
    - Remove extra whitespace
    - Remove common punctuation
    """
    if not isinstance(name, str):
        return ""

    # Convert to lowercase
    name = name.lower()

    # Remove château/chateau variations and common prefixes
    name = re.sub(r'\bch[âa]teau\b', '', name)
    name = re.sub(r'\bdomaine\b', '', name)
    name = re.sub(r'\bch[âa]teau\b', '', name)

    # Remove special characters but keep spaces
    name = re.sub(r'[^\w\s]', ' ', name)

    # Remove extra whitespace
    name = ' '.join(name.split())

    return name.strip()


def calculate_similarity(text1, text2):
    """
    Calculate similarity ratio between two strings (0-1).
    Uses both full string matching and partial matching for better results.
    """
    text1_norm = normalize_wine_name(text1)
    text2_norm = normalize_wine_name(text2)

    if not text1_norm or not text2_norm:
        return 0.0

    # Full string similarity
    full_similarity = SequenceMatcher(None, text1_norm, text2_norm).ratio()

    # Check if one is contained in the other (partial match)
    if text1_norm in text2_norm or text2_norm in text1_norm:
        # Boost score for substring matches
        full_similarity = max(full_similarity, 0.7)

    # Word-level matching (check if key words are shared)
    words1 = set(text1_norm.split())
    words2 = set(text2_norm.split())

    # Remove common filler words
    filler_words = {'the', 'de', 'di', 'du', 'della', 'des', 'le', 'la', 'del'}
    words1 = words1 - filler_words
    words2 = words2 - filler_words

    if words1 and words2:
        # Jaccard similarity for word sets
        word_overlap = len(words1 & words2) / len(words1 | words2)
        # Combine full string and word-level similarity
        combined_similarity = max(full_similarity, word_overlap * 0.9)
    else:
        combined_similarity = full_similarity

    return combined_similarity


def detect_market_price_context(text, price_match_start):
    """
    Detect if a price is a market price reference (should use 1.08 conversion).
    Returns True if market price indicators are found VERY CLOSE to this specific price.
    """
    # Look in TIGHT context (only 40 chars before, 30 after) to avoid false positives
    context_before = text[max(0, price_match_start - 40):price_match_start].lower()
    context_after = text[price_match_start:min(len(text), price_match_start + 30)].lower()

    # Market price indicators that must be IMMEDIATELY before "CHF"
    # Pattern: "market price stands at approximately CHF 100"
    #          "swiss market price ... CHF 100"
    immediate_before_patterns = [
        r'market\s+price\s+chf\s*$',           # "market price CHF" directly before number
        r'approximately\s+chf\s*$',             # "approximately CHF" directly before
        r'around\s+chf\s*$',                    # "around CHF" directly before
        r'stands\s+at.*chf\s*$',                # "stands at ... CHF" directly before
    ]

    # Check if "CHF" appears right before the price position in original text
    # This helps identify if this is part of "CHF XX" pattern
    chf_before_price = text[max(0, price_match_start - 5):price_match_start]
    has_chf_prefix = bool(re.search(r'chf\s*$', chf_before_price, re.IGNORECASE))

    # Only apply market price detection if CHF prefix exists
    if has_chf_prefix:
        for pattern in immediate_before_patterns:
            if re.search(pattern, context_before):
                return True

    # Market price indicators AFTER the price (in parentheses)
    # Pattern: "100.00 EUR + VAT (market price EUR 108.00)"
    market_indicators_after = [
        r'^\s*\)?\s*\(?\s*market\s+price',     # "(market price" right after
    ]

    for indicator in market_indicators_after:
        if re.search(indicator, context_after):
            return True

    return False


def detect_quantity_indicator(text, price_match_start):
    """
    Detect quantity indicators near the price (e.g., "36x", "36 bottles").
    Returns the minimum quantity (0 or 36) based on context.
    IMPORTANT: Must be VERY close to the price to avoid false positives.
    """
    # Look in TIGHT context around the price
    context_before = text[max(0, price_match_start - 30):price_match_start].lower()
    context_after = text[price_match_start:min(len(text), price_match_start + 35)].lower()

    # Check for 36-bottle indicators (MUST BE VERY CLOSE)
    # Pattern 1: "36x" directly before the price (within 10 chars)
    if re.search(r'36\s*x\b', context_before[-10:]):
        return 36

    # Pattern 2: "36+ bottle" directly before or after (within 25 chars)
    if re.search(r'36\s*\+\s*bottle', context_before):
        return 36
    if re.search(r'^[^a-z]*36\s*\+\s*bottle', context_after):
        return 36

    # Pattern 3: "if you take 36 bottles" AFTER the price (not before!)
    # This must be in the context_after, not before
    if re.search(r'if\s+you\s+take\s+36\s+bottle', context_after):
        return 36

    # Pattern 4: Price at end of "36 bottles" phrase (within 20 chars before)
    if re.search(r'36\s+bottles?\s*(at|for)?\s*chf', context_before):
        return 36

    # Default to 0 (normal price, no minimum quantity)
    return 0


def detect_size_indicator(text, price_match_start):
    """
    Detect size indicators near the price (e.g., "Magnum" for 150cl).
    Returns the size in cl (75 for standard, 150 for Magnum, etc.)
    """
    # Look in context around the price (within 50 chars before)
    context_before = text[max(0, price_match_start - 50):price_match_start].lower()

    # Check for Magnum indicator (150cl)
    if re.search(r'\bmagnum\b', context_before):
        return 150.0

    # Default to standard bottle size (75cl)
    return 75.0


def extract_vintage_from_context(text, price_match_start):
    """
    Extract vintage year from context around the price.
    Returns integer year or None.
    """
    # Look in wider context (up to 600 chars before to catch vintage at paragraph start)
    context = text[max(0, price_match_start - 600):min(len(text), price_match_start + 100)]

    # Find 4-digit years (vintage years typically 1990-2030)
    year_matches = re.findall(r'\b(19[9]\d|20[0-3]\d)\b', context)

    if year_matches:
        # Return the most recent/last mentioned year
        try:
            return int(year_matches[-1])
        except:
            return None

    return None


def extract_wine_name_from_context(text, price_match_start):
    """
    Extract potential wine name from context around the price.
    Looks backwards from the price position to find wine name.
    """
    # Get text before the price (up to 400 characters for better context)
    context_before = text[max(0, price_match_start - 400):price_match_start]

    wine_candidates = []

    # Pattern 1: FIRST colon in the context (usually the wine name at paragraph start)
    # e.g., "Château Rieussec 2019: ... price"
    # Find ALL colons, prefer the FIRST one (which is usually the wine name)
    all_colons = list(re.finditer(r'([A-ZÀ-ÿ][^\n:]{3,60})[:]\s*', context_before))
    if all_colons:
        # Take the FIRST colon match (likely the wine name)
        first_colon = all_colons[0]
        candidate = first_colon.group(1).strip()
        # Remove common trailing words
        candidate = re.sub(r'\s+(at|from|for|with|the|a|an)$', '', candidate, flags=re.IGNORECASE)
        wine_candidates.append(candidate)

        # Also consider the LAST colon if different (might be more specific context)
        if len(all_colons) > 1:
            last_colon = all_colons[-1]
            candidate_last = last_colon.group(1).strip()
            candidate_last = re.sub(r'\s+(at|from|for|with|the|a|an)$', '', candidate_last, flags=re.IGNORECASE)
            # Only add if it's different and doesn't contain "CHF" or price-related words
            if candidate_last != candidate and not re.search(r'chf|price|offer', candidate_last, re.IGNORECASE):
                wine_candidates.append(candidate_last)

    # Pattern 2: Quoted text (wine names often in quotes)
    # e.g., "the famous "Château Pavie" at"
    quote_matches = re.findall(r'["""]([^"""]{3,60})["""]', context_before)
    for match in quote_matches:
        # Prefer quotes with capitalized content
        if re.search(r'[A-ZÀ-ÿ]', match):
            wine_candidates.append(match.strip())

    # Pattern 3: Château/Domaine followed by name
    # e.g., "Château Montrose 2021"
    chateau_pattern = re.findall(r'\b([CcDd]h[âa]teau|Domaine|Dom\.)\s+([A-ZÀ-ÿ][^\n:,.]{3,40})', context_before)
    for prefix, name in chateau_pattern:
        wine_candidates.append(f"{prefix} {name}".strip())

    # Pattern 4: Producer name patterns (e.g., "Penfolds 2019", "Aalto 2023")
    # Match: Capitalized word(s) optionally followed by year
    producer_pattern = re.findall(
        r'\b([A-ZÀ-ÿ][a-zà-ÿ]+(?:\s+[A-ZÀ-ÿ][a-zà-ÿ]+){0,3})\s+(?:\d{4})?',
        context_before
    )
    if producer_pattern:
        # Get last few capitalized phrases
        wine_candidates.extend(producer_pattern[-3:])

    # Pattern 5: Text between line start and dash/colon
    # e.g., "Aalto 2023: ..." or "Dominus 2016: ..."
    line_start = context_before.split('\n')[-1] if '\n' in context_before else context_before
    line_pattern = re.match(r'^([A-ZÀ-ÿ][^\n:–-]{3,60}?)[:–-]', line_start.strip())
    if line_pattern:
        wine_candidates.append(line_pattern.group(1).strip())

    # Filter and clean candidates
    cleaned_candidates = []
    for candidate in wine_candidates:
        # Remove year patterns at the end
        candidate = re.sub(r'\s+\d{4}\s*$', '', candidate)
        # Remove common noise words at the end
        candidate = re.sub(r'\s+(at|from|for|with|price|the|a|an)$', '', candidate, flags=re.IGNORECASE)
        # Remove extra whitespace
        candidate = ' '.join(candidate.split())
        if len(candidate) >= 3:
            cleaned_candidates.append(candidate)

    # Return the longest and most specific candidate
    if cleaned_candidates:
        # Prefer longer candidates that likely contain more specific info
        return max(cleaned_candidates, key=lambda x: (len(x), x.count(' ')))

    return ""


def load_data_and_document():
    """Loads the Excel data, creates the conversion map, and loads the Word document."""
    conversion_map = {}
    wine_data_map = defaultdict(list)  # CHF -> list of full row data
    item_number_map = {}  # Item No. -> wine data (bulletproof matching)
    duplicate_chf_prices = set()
    doc = None
    df_full = None  # Store full dataframe for reference

    # 1. Load Excel File and Create Conversion Dictionary
    try:
        df = pd.read_excel(EXCEL_FILE_PATH)
        df_full = df.copy()

        # Standardize columns
        df['CHF_KEY_FORMATTED'] = df[CHF_COL].astype(float).round(2).apply(lambda x: f'{x:.2f}')
        df['EUR_VALUE_FORMATTED'] = df[EUR_COL].astype(float).round(2).apply(lambda x: f'{x:.2f}')
        df['WINE_NAME_NORMALIZED'] = df[WINE_NAME_COL].astype(str)

        # Find duplicate CHF prices (same CHF, different EUR)
        chf_eur_mapping = df.groupby('CHF_KEY_FORMATTED')['EUR_VALUE_FORMATTED'].apply(set).to_dict()
        for chf, eur_set in chf_eur_mapping.items():
            if len(eur_set) > 1:
                duplicate_chf_prices.add(chf)

        # Create wine data mapping with all relevant columns
        for _, row in df.iterrows():
            chf = row['CHF_KEY_FORMATTED']
            eur = row['EUR_VALUE_FORMATTED']
            wine = row['WINE_NAME_NORMALIZED']

            # Extract additional columns for filtering
            campaign_subtype = str(row.get(CAMPAIGN_SUBTYPE_COL, '')).strip().lower()
            campaign_type = str(row.get(CAMPAIGN_TYPE_COL, '')).strip().lower()
            size = row.get(SIZE_COL, 0)
            min_qty = row.get(MIN_QUANTITY_COL, 0)
            competitor_code = row.get(COMPETITOR_CODE_COL, None)
            producer_name = str(row.get(PRODUCER_NAME_COL, '')).strip()
            vintage_raw = row.get(VINTAGE_COL, None)
            # Convert vintage to int for proper comparison with context_vintage (which is int)
            try:
                vintage = int(vintage_raw) if pd.notna(vintage_raw) else None
            except (ValueError, TypeError):
                vintage = None
            item_no = row.get(ITEM_NO_COL, None)

            wine_data = {
                'wine_name': wine,
                'eur_value': eur,
                'campaign_subtype': campaign_subtype,
                'campaign_type': campaign_type,
                'size': size,
                'min_quantity': min_qty,
                'competitor_code': competitor_code,
                'producer_name': producer_name,
                'vintage': vintage,
                'item_no': item_no,
                'chf_value': chf
            }

            wine_data_map[chf].append(wine_data)

            # Build item_number_map for bulletproof matching
            # Item No. is unique per wine+vintage+size (same for qty=0 and qty=36)
            if pd.notna(item_no):
                try:
                    item_key = int(float(item_no))  # Handle both int and float strings
                    if item_key not in item_number_map:
                        item_number_map[item_key] = []
                    item_number_map[item_key].append(wine_data)
                except (ValueError, TypeError):
                    # Skip non-numeric Item Numbers (e.g., "ACCESSORIES")
                    pass

        # For non-duplicate prices, create simple conversion map
        for chf, eur_set in chf_eur_mapping.items():
            if len(eur_set) == 1:
                conversion_map[chf] = list(eur_set)[0]

        print("✅ Excel file loaded successfully.")
        print(f"   - Found {len(conversion_map)} unique CHF->EUR conversions")
        print(f"   - Found {len(duplicate_chf_prices)} duplicate CHF prices requiring wine name matching")
        print(f"   - Loaded {len(df)} rows with full metadata (Campaign Type, Size, Min Quantity)")
        print(f"   - Built Item No. map with {len(item_number_map)} unique items for bulletproof matching")

    except FileNotFoundError:
        print(f"❌ Error: Excel file not found at {EXCEL_FILE_PATH}")
    except KeyError as e:
        print(f"❌ Error: Column not found. Details: {e}")
    except Exception as e:
        print(f"❌ Unexpected error loading Excel file: {e}")

    # 2. Load Word Document
    try:
        doc = Document(WORD_FILE_PATH)
        print(f"✅ Word file loaded from {WORD_FILE_PATH}")
    except FileNotFoundError:
        print(f"❌ Error: Word file not found at {WORD_FILE_PATH}")
    except Exception as e:
        print(f"❌ Unexpected error loading Word file: {e}")

    return doc, conversion_map, wine_data_map, duplicate_chf_prices, item_number_map


def round_to_5_or_0(value):
    """
    Round value up to the nearest number ending in 5 or 0.
    Used for prices above 300 CHF.
    Examples: 1162 -> 1165, 1163 -> 1165, 1167 -> 1170
    """
    # Get the last digit
    last_digit = int(value) % 10

    if last_digit == 0 or last_digit == 5:
        return value  # Already ends in 0 or 5
    elif last_digit < 5:
        # Round up to 5 (e.g., 1162 -> 1165)
        return int(value) + (5 - last_digit)
    else:
        # Round up to next 10 (e.g., 1167 -> 1170)
        return int(value) + (10 - last_digit)


def find_best_wine_match(chf_price, context_wine_name, wine_data_map, detected_quantity=0,
                         context_vintage=None, context_producer=None, detected_size=75.0, item_number_map=None):
    """
    Find the best EUR conversion for a CHF price using Item No. matching (bulletproof) and wine name matching.

    Args:
        chf_price: CHF price string (e.g., "42.00")
        context_wine_name: Wine name extracted from context
        wine_data_map: Dictionary mapping CHF to list of wine data dicts
        detected_quantity: Detected minimum quantity (0 or 36)
        context_vintage: Vintage year extracted from context (optional)
        context_producer: Producer name extracted from context (optional)
        detected_size: Detected size in cl (75 for standard, 150 for Magnum)
        item_number_map: Dictionary mapping Item No. to wine data (for bulletproof matching)

    Returns (eur_value, match_quality, wine_data_option)
    """
    if chf_price not in wine_data_map:
        return None, 'not_found', None

    wine_options = wine_data_map[chf_price]

    # BULLETPROOF MATCHING: Try Item No. matching FIRST if we have vintage
    # Item No. is unique per wine+vintage+size (same for both qty=0 and qty=36)
    # This works for BOTH single and multiple options - Item No. is always most reliable
    if item_number_map and context_vintage:
        # Check all options for this CHF price
        for option in wine_options:
            item_no = option.get('item_no')
            if pd.notna(item_no):
                try:
                    item_key = int(float(item_no))  # Handle both int and float strings
                except (ValueError, TypeError):
                    # Skip non-numeric Item Numbers
                    continue
                if item_key in item_number_map:
                    # Get all entries for this Item No. (should be qty=0 and qty=36 variants)
                    item_entries = item_number_map[item_key]
                    # Check if vintage and size match
                    for entry in item_entries:
                        if (entry.get('vintage') == context_vintage and
                            entry.get('size') == detected_size and
                            entry.get('min_quantity') == detected_quantity and
                            entry.get('chf_value') == chf_price):
                            # BULLETPROOF MATCH!
                            return entry['eur_value'], 'item_no_match', entry

    # Fallback: If only one option and Item No. didn't match, use it
    if len(wine_options) == 1:
        return wine_options[0]['eur_value'], 'exact', wine_options[0]

    # Continue with existing matching logic if Item No. match not found

    # Multiple options - apply filters first
    # Filter 1: Campaign Sub-Type = "Normal"
    filtered_options = [opt for opt in wine_options if opt['campaign_subtype'] == 'normal']

    if not filtered_options:
        filtered_options = wine_options

    # Filter 2: Size = detected_size (75 for standard, 150 for Magnum, etc.)
    size_matched_options = [opt for opt in filtered_options if opt['size'] == detected_size]

    if size_matched_options:
        filtered_options = size_matched_options

    # Filter 3: Match quantity (0 for normal, 36 for bulk)
    quantity_matched_options = [opt for opt in filtered_options if opt['min_quantity'] == detected_quantity]

    if quantity_matched_options:
        filtered_options = quantity_matched_options

        # VALIDATION: If detected_quantity == 36, ensure the price is different from normal (qty=0)
        # This prevents matching a normal price when "36 bottles" is mentioned
        # IMPORTANT: Only compare prices for the SAME wine to avoid false filtering
        if detected_quantity == 36 and len(filtered_options) > 1:
            # For each 36-bottle option, check if there's a matching normal (qty=0) option
            # for the SAME wine with the SAME EUR price
            validated_options = []
            for opt_36 in filtered_options:
                # Find normal prices for the SAME wine
                same_wine_normal = [opt for opt in wine_options
                                  if opt['min_quantity'] == 0
                                  and opt['size'] == detected_size
                                  and opt['wine_name'] == opt_36['wine_name']]

                if same_wine_normal:
                    # Check if this 36x price differs from normal price for same wine
                    normal_eur_values = set(opt['eur_value'] for opt in same_wine_normal)
                    if opt_36['eur_value'] not in normal_eur_values:
                        # This 36x price is different from normal, keep it
                        validated_options.append(opt_36)
                else:
                    # No normal price for this wine, keep the 36x option
                    validated_options.append(opt_36)

            if validated_options:
                filtered_options = validated_options

    # Filter 4: Competitor Code is empty (NaN) - PREFER these rows
    empty_competitor_options = [opt for opt in filtered_options
                                if pd.isna(opt['competitor_code']) or opt['competitor_code'] == '']

    if empty_competitor_options:
        filtered_options = empty_competitor_options

    # Filter 5: Campaign Type = "PRIVATE" (only for quantity=0)
    if len(filtered_options) > 1 and detected_quantity == 0:
        private_options = [opt for opt in filtered_options if opt['campaign_type'] == 'private']
        if private_options:
            filtered_options = private_options

    # Filter 6: Match vintage if available
    if context_vintage and len(filtered_options) > 1:
        vintage_matched = [opt for opt in filtered_options
                          if opt['vintage'] == context_vintage
                          or opt['vintage'] == float(context_vintage)
                          or opt['vintage'] == str(context_vintage)]
        if vintage_matched:
            filtered_options = vintage_matched

    # NEW: Filter 7: Prefer LAST row (bottom-most) when multiple Size=75 matches remain
    # This helps when there are multiple identical wines with different campaigns
    if len(filtered_options) > 1:
        # Add row index to each option (stored during data loading)
        # For now, we'll use the last item in the list as it represents the bottom-most row
        # since pandas iterates top-to-bottom and we append in order
        filtered_options = [filtered_options[-1]]

    # After filtering, if only one option remains, use it
    if len(filtered_options) == 1:
        return filtered_options[0]['eur_value'], 'fuzzy_filtered', filtered_options[0]

    # Multiple options remain - use intelligent matching
    if not context_wine_name and not context_producer:
        # Use price proximity (choose EUR closest to CHF * 1.08)
        if filtered_options:
            chf_float = float(chf_price)
            expected_eur = chf_float * 1.08
            best_option = min(filtered_options, key=lambda opt: abs(float(opt['eur_value']) - expected_eur))
            return best_option['eur_value'], 'price_proximity', best_option
        return None, 'ambiguous', None

    # Calculate similarity scores for each option
    scored_options = []
    for option in filtered_options:
        score = 0.0

        # Wine name similarity
        if context_wine_name:
            wine_similarity = calculate_similarity(context_wine_name, option['wine_name'])
            score += wine_similarity * 2.0  # Weight: 2.0

        # Producer name similarity
        if context_producer and option['producer_name']:
            producer_similarity = calculate_similarity(context_producer, option['producer_name'])
            score += producer_similarity * 1.5  # Weight: 1.5

        # Price proximity bonus (prefer EUR closest to CHF * 1.08)
        chf_float = float(chf_price)
        expected_eur = chf_float * 1.08
        eur_float = float(option['eur_value'])
        price_diff = abs(eur_float - expected_eur)
        # Normalize: closer = higher score
        price_proximity_score = max(0, 1.0 - (price_diff / expected_eur))
        score += price_proximity_score * 0.5  # Weight: 0.5

        scored_options.append((option, score))

    # Sort by score
    scored_options.sort(key=lambda x: x[1], reverse=True)
    best_option, best_score = scored_options[0]

    # If good match (score > threshold), use it
    if best_score >= 1.0:  # Threshold for combined score
        return best_option['eur_value'], 'fuzzy', best_option

    # Otherwise, use price proximity
    chf_float = float(chf_price)
    expected_eur = chf_float * 1.08
    best_by_price = min(filtered_options, key=lambda opt: abs(float(opt['eur_value']) - expected_eur))
    return best_by_price['eur_value'], 'price_proximity', best_by_price


def export_to_lines_excel(conversion_records, output_path):
    """
    Export conversion records to Lines.xlsx file.

    Args:
        conversion_records: List of conversion record dicts
        output_path: Path to save the Lines.xlsx file
    """
    # Prepare data for Excel export
    lines_data = []

    # Process all records and extract data
    processed_records = []
    for record in conversion_records:
        wine_data = record.get('wine_data')
        match_type = record.get('match_type', '')

        # PRIORITY 1: Use exact data from Excel wine_data when available
        if wine_data and isinstance(wine_data, dict):
            wine_name = wine_data.get('wine_name', '')
            vintage = wine_data.get('vintage', '')
            size = wine_data.get('size', '')
            producer_name = wine_data.get('producer_name', '')
            min_quantity = wine_data.get('min_quantity', 0)
            item_no = wine_data.get('item_no', '')
        # PRIORITY 2: Fallback to context extraction for formula-calculated prices
        else:
            wine_name = record.get('context_wine_name', '')
            vintage = record.get('context_vintage', '')
            size = record.get('detected_size', '')
            producer_name = record.get('context_producer', '')
            min_quantity = record.get('detected_min_quantity', 0)
            item_no = ''

        processed_records.append({
            'wine_name': wine_name if wine_name else '',
            'vintage': vintage if vintage else '',
            'size': size if size else '',
            'producer_name': producer_name if producer_name else '',
            'min_quantity': min_quantity if min_quantity is not None else 0,
            'chf_price': record.get('chf_price', ''),
            'eur_price': record.get('eur_price', ''),
            'match_type': match_type,
            'item_no': item_no if item_no else ''
        })

    # Deduplicate: Keep only 2 rows per wine (Min Qty 0 and 36)
    # Group by wine identifier (wine_name + vintage + producer_name)
    wine_groups = {}
    for rec in processed_records:
        # Create unique key for each wine
        wine_key = (rec['wine_name'], rec['vintage'], rec['producer_name'], rec['size'])

        if wine_key not in wine_groups:
            wine_groups[wine_key] = {'qty_0': None, 'qty_36': None}

        min_qty = rec['min_quantity']

        # Keep best match for each quantity level
        if min_qty == 0:
            # For qty=0, keep if we don't have one yet, or if this is a better match
            if wine_groups[wine_key]['qty_0'] is None:
                wine_groups[wine_key]['qty_0'] = rec
            elif rec['match_type'] in ['direct', 'fuzzy_filtered', 'exact_match']:
                # Prefer Excel matches over formula calculations
                if wine_groups[wine_key]['qty_0']['match_type'] in ['fallback_1.08', 'market_price_1.08']:
                    wine_groups[wine_key]['qty_0'] = rec
        elif min_qty == 36:
            # For qty=36, keep if we don't have one yet, or if this is a better match
            if wine_groups[wine_key]['qty_36'] is None:
                wine_groups[wine_key]['qty_36'] = rec
            elif rec['match_type'] in ['direct', 'fuzzy_filtered', 'exact_match']:
                # Prefer Excel matches over formula calculations
                if wine_groups[wine_key]['qty_36']['match_type'] in ['fallback_1.08', 'market_price_1.08']:
                    wine_groups[wine_key]['qty_36'] = rec

    # Build final lines data with only kept records
    group_code = 1
    for wine_key, quantities in wine_groups.items():
        # Add qty=0 row if exists
        if quantities['qty_0']:
            rec = quantities['qty_0']
            row_data = {
                'Wine Name': rec['wine_name'],
                'Vintage Code': rec['vintage'],
                'Size': rec['size'],
                'Producer Name': rec['producer_name'],
                'Minimum Quantity': rec['min_quantity'],
                'Unit Price': rec['chf_price'],
                'Unit Price Incl. VAT': '',
                'Unit Price (EUR)': rec['eur_price'],
                'Main Offer Comment': '',
                'Competitor Code': '',
                'Group Code': group_code,
                'Match Type': rec['match_type'],
                'Item No.': rec['item_no']
            }
            lines_data.append(row_data)

        # Add qty=36 row if exists
        if quantities['qty_36']:
            rec = quantities['qty_36']
            row_data = {
                'Wine Name': rec['wine_name'],
                'Vintage Code': rec['vintage'],
                'Size': rec['size'],
                'Producer Name': rec['producer_name'],
                'Minimum Quantity': rec['min_quantity'],
                'Unit Price': rec['chf_price'],
                'Unit Price Incl. VAT': '',
                'Unit Price (EUR)': rec['eur_price'],
                'Main Offer Comment': '',
                'Competitor Code': '',
                'Group Code': group_code,
                'Match Type': rec['match_type'],
                'Item No.': rec['item_no']
            }
            lines_data.append(row_data)

        group_code += 1  # Increment group code per wine

    # Create DataFrame and export to Excel
    df_lines = pd.DataFrame(lines_data)

    # Save to Excel (overwrite if exists)
    df_lines.to_excel(output_path, index=False, sheet_name='Lines')

    print(f"\n✅ Lines Excel file saved as: {output_path}")
    print(f"   Total lines exported: {len(lines_data)}")
    print(f"   Unique wines: {len(wine_groups)}")


def replace_and_highlight(paragraph, conversion_map, wine_data_map, duplicate_chf_prices,
                         all_numbers_found, conversion_stats, conversion_records, item_number_map):
    """
    Performs search and replace at the run level, applying highlighting based on conversion type.
    Also tracks conversion data for Excel export.

    Args:
        conversion_records: List to append conversion record dicts to
        item_number_map: Dictionary mapping Item No. to wine data (for bulletproof matching)
    """

    local_replacements = 0

    # 1. Replace CHF with EUR in the paragraph text
    original_text = paragraph.text

    text = original_text

    # Find all number matches with their positions (three patterns)
    # Pattern 1: Standard XX.XX format (including Swiss format with apostrophe like 1'500.00 or 1'500.00)
    matches_with_positions = []
    for m in re.finditer(NUMBER_PATTERN, text):
        # Remove both regular apostrophe (') and curly quote (') for Excel lookup
        number_clean = m.group().replace("'", "").replace("'", "")
        matches_with_positions.append((number_clean, m.start()))

    # Pattern 2: "CHF 100" style (convert to "100.00" format)
    chf_no_decimal_matches = re.finditer(CHF_NUMBER_NO_DECIMAL, text)
    for match in chf_no_decimal_matches:
        # Remove both regular apostrophe and curly quote
        number = match.group(1).replace("'", "").replace("'", "")
        formatted_number = f"{number}.00"
        # Add to matches with the position of the number (not CHF)
        matches_with_positions.append((formatted_number, match.start() + match.group().find(match.group(1))))

    # Pattern 3: "190 CHF" style (NUMBER THEN CHF, without decimals)
    number_then_chf_matches = re.finditer(NUMBER_THEN_CHF, text)
    for match in number_then_chf_matches:
        # Remove both regular apostrophe and curly quote
        number = match.group(1).replace("'", "").replace("'", "")
        # Check if this number already has a decimal match (skip if so)
        if not any(chf == f"{number}.00" or chf == number for chf, _ in matches_with_positions):
            formatted_number = f"{number}.00"
            matches_with_positions.append((formatted_number, match.start()))

    # Pattern 4: "33.00CHF" style (NUMBER.XX directly followed by CHF, no space)
    number_nospace_chf_matches = re.finditer(NUMBER_NOSPACE_CHF, text)
    for match in number_nospace_chf_matches:
        # Remove both regular apostrophe and curly quote
        number = match.group(1).replace("'", "").replace("'", "")
        # Check if already matched (avoid duplicates)
        if not any(chf == number for chf, _ in matches_with_positions):
            matches_with_positions.append((number, match.start()))

    # Pattern 5: "Magnum XX.XX" style (number after Magnum keyword, no CHF indicator)
    # This is a special case for Magnum bottles where CHF may be omitted
    magnum_pattern = r"\bMagnum\s+(\d+(?:['\u2019]\d{3})*\.\d{2})"
    magnum_matches = re.finditer(magnum_pattern, text, re.IGNORECASE)
    for match in magnum_matches:
        number = match.group(1).replace("'", "").replace("'", "")
        # Check if already matched (avoid duplicates)
        if not any(chf == number for chf, _ in matches_with_positions):
            matches_with_positions.append((number, match.start() + match.group().find(match.group(1))))

    # Pattern 6: "36x XX.XX" style (number after 36x, may or may not have CHF)
    # Special case for 36-bottle pricing where CHF may be omitted or attached
    # This pattern should NOT match if Pattern 4 already matched (avoid "36x 33.00CHF")
    thirtysix_pattern = r"\b36\s*x\s+(\d+(?:['\u2019]\d{3})*\.\d{2})(?![Cc][Hh][Ff])"
    thirtysix_matches = re.finditer(thirtysix_pattern, text, re.IGNORECASE)
    for match in thirtysix_matches:
        number = match.group(1).replace("'", "").replace("'", "")
        # Check if already matched (avoid duplicates)
        if not any(chf == number for chf, _ in matches_with_positions):
            matches_with_positions.append((number, match.start() + match.group().find(match.group(1))))

    # Dictionary of replacements: {chf_str: (new_eur_str, highlight_color_index, context)}
    replacements_to_do = {}

    for chf_str, position in matches_with_positions:

        # --- VINTAGE/YEAR SKIP CHECK ---
        # Skip numbers that look like years (1000-9999) UNLESS they have CHF/EUR context
        try:
            float_val = float(chf_str)
            if float_val.is_integer() and 1000 <= int(float_val) <= 9999:
                # Check if CHF or EUR appears within 20 characters of this position
                context_start = max(0, position - 20)
                context_end = min(len(text), position + len(chf_str) + 20)
                context = text[context_start:context_end]

                # If CHF or EUR is in the context, this is a price, not a year
                if re.search(r'\b(CHF|EUR|chf|eur)\b', context):
                    pass  # Don't skip, it's a price
                else:
                    continue  # Skip, it's likely a year
        except ValueError:
            continue

        all_numbers_found.append(chf_str)

        eur_value = None
        highlight_color = None
        match_type = None

        # Detect market price context (should force 1.08 conversion)
        is_market_price = detect_market_price_context(text, position)

        # Extract context information
        context_wine = extract_wine_name_from_context(text, position)
        context_vintage = extract_vintage_from_context(text, position)

        # Try to extract producer name from wine name
        context_producer = None
        if context_wine:
            # Check if wine name contains producer patterns
            producer_match = re.search(r'(Aalto|Colgin|Roederer|Louis Roederer)', context_wine, re.IGNORECASE)
            if producer_match:
                context_producer = producer_match.group(1)

        # Detect quantity indicator (0 or 36)
        detected_quantity = detect_quantity_indicator(text, position)

        # Detect size indicator (75 or 150 for Magnum)
        detected_size = detect_size_indicator(text, position)

        # If it's a market price reference, always use 1.08 conversion with RED highlight
        if is_market_price:
            try:
                raw_eur = float(chf_str) * 1.08
                floored_eur = math.floor(raw_eur)
                eur_value = f"{floored_eur:.2f}"
                highlight_color = WD_COLOR_INDEX.RED
                match_type = 'market_price_1.08'
                conversion_stats['fallback'] += 1
            except ValueError:
                continue

        # Check if it's a known non-duplicate price
        elif chf_str in conversion_map:
            # Try bulletproof Item No. matching first, even for direct matches
            wine_data = None

            if context_vintage and item_number_map and chf_str in wine_data_map:
                # Try Item No. matching
                eur_from_item, quality_item, wine_from_item = find_best_wine_match(
                    chf_str, context_wine, wine_data_map, detected_quantity,
                    context_vintage, context_producer, detected_size, item_number_map
                )

                if quality_item == 'item_no_match':
                    # Bulletproof match found!
                    eur_value = eur_from_item
                    wine_data = wine_from_item
                    match_type = 'item_no_match'
                    conversion_stats['exact_matched'] += 1
                else:
                    # No Item No. match, use direct conversion
                    eur_value = conversion_map[chf_str]
                    match_type = 'direct'
                    conversion_stats['direct'] += 1

                    # Get wine data for tracking
                    if chf_str in wine_data_map:
                        options = wine_data_map[chf_str]
                        if len(options) == 1:
                            wine_data = options[0]
                        else:
                            filtered = [opt for opt in options
                                      if opt.get('size') == detected_size
                                      and opt.get('min_quantity') == detected_quantity]
                            wine_data = filtered[0] if filtered else options[0]
            else:
                # No vintage or Item No. map, use direct conversion
                eur_value = conversion_map[chf_str]
                match_type = 'direct'
                conversion_stats['direct'] += 1

                # Get wine data for tracking
                if chf_str in wine_data_map:
                    options = wine_data_map[chf_str]
                    if len(options) == 1:
                        wine_data = options[0]
                    else:
                        filtered = [opt for opt in options
                                  if opt.get('size') == detected_size
                                  and opt.get('min_quantity') == detected_quantity]
                        wine_data = filtered[0] if filtered else options[0]

            # Apply rounding rule for prices above 300 CHF
            chf_float = float(chf_str)
            if chf_float > 300:
                eur_float = float(eur_value)
                rounded_eur = round_to_5_or_0(eur_float)
                eur_value = f"{rounded_eur:.2f}"

        # Check if it's a duplicate requiring wine name matching
        elif chf_str in duplicate_chf_prices:
            eur_value, quality, wine_data = find_best_wine_match(
                chf_str, context_wine, wine_data_map, detected_quantity,
                context_vintage, context_producer, detected_size, item_number_map
            )

            if quality == 'item_no_match':
                # BULLETPROOF Item No. match - no highlighting needed
                # Apply rounding rule for prices above 300 CHF
                chf_float = float(chf_str)
                if chf_float > 300 and eur_value:
                    eur_float = float(eur_value)
                    rounded_eur = round_to_5_or_0(eur_float)
                    eur_value = f"{rounded_eur:.2f}"

                match_type = 'item_no_match'
                conversion_stats['exact_matched'] += 1
            elif quality == 'fuzzy' or quality == 'fuzzy_filtered' or quality == 'price_proximity':
                # Apply rounding rule for prices above 300 CHF
                chf_float = float(chf_str)
                if chf_float > 300 and eur_value:
                    eur_float = float(eur_value)
                    rounded_eur = round_to_5_or_0(eur_float)
                    eur_value = f"{rounded_eur:.2f}"

                highlight_color = WD_COLOR_INDEX.GREEN
                match_type = quality
                conversion_stats['fuzzy_matched'] += 1
            elif quality == 'exact':
                # Apply rounding rule for prices above 300 CHF
                chf_float = float(chf_str)
                if chf_float > 300 and eur_value:
                    eur_float = float(eur_value)
                    rounded_eur = round_to_5_or_0(eur_float)
                    eur_value = f"{rounded_eur:.2f}"

                match_type = 'exact_match'
                conversion_stats['exact_matched'] += 1
            elif quality == 'ambiguous':
                # Apply 1.08 fallback and highlight yellow
                try:
                    raw_eur = float(chf_str) * 1.08
                    floored_eur = math.floor(raw_eur)

                    # Apply rounding rule for prices above 300 CHF
                    chf_float = float(chf_str)
                    if chf_float > 300:
                        floored_eur = round_to_5_or_0(floored_eur)

                    eur_value = f"{floored_eur:.2f}"
                    highlight_color = WD_COLOR_INDEX.YELLOW
                    match_type = 'ambiguous'
                    conversion_stats['ambiguous'] += 1
                except ValueError:
                    continue

        # Not in Excel at all - use 1.08 fallback
        else:
            try:
                raw_eur = float(chf_str) * 1.08
                floored_eur = math.floor(raw_eur)

                # Apply rounding rule for prices above 300 CHF
                chf_float = float(chf_str)
                if chf_float > 300:
                    floored_eur = round_to_5_or_0(floored_eur)

                eur_value = f"{floored_eur:.2f}"
                highlight_color = WD_COLOR_INDEX.RED
                match_type = 'fallback_1.08'
                conversion_stats['fallback'] += 1
            except ValueError:
                continue

        if eur_value:
            replacements_to_do[chf_str] = (eur_value, highlight_color, context_wine, match_type)

            # Track conversion for Excel export
            # Store both context and wine_data - export function will prioritize wine_data
            conversion_record = {
                'chf_price': chf_str,
                'eur_price': eur_value,
                'context_wine_name': context_wine if context_wine else '',
                'context_vintage': context_vintage if context_vintage else '',
                'context_producer': context_producer if context_producer else '',
                'detected_size': detected_size,
                'detected_min_quantity': detected_quantity,
                'match_type': match_type,
                'wine_data': wine_data if 'wine_data' in locals() else None
            }
            conversion_records.append(conversion_record)

    # 2. Perform number replacement with highlighting (rebuilding runs)
    new_text = text

    # Create a list of all replacements with their exact positions
    # Format: (start_pos, end_pos, replacement_text)
    all_replacements = []

    # We need to find each CHF price in the original text and replace it
    # The key is to match the EXACT format that appears in the text
    for chf_str, (eur_value, _, _, _) in replacements_to_do.items():
        # chf_str is the normalized value (e.g., "1150.00" or "190.00")
        # We need to find it in text which might be "1'150.00 CHF" or "190 CHF"

        chf_int = chf_str.replace('.00', '')  # "1150" or "190"

        # Build patterns to find ALL possible representations
        # NOTE: Use ['\u2019]? to match both regular apostrophe (') and curly quote (')
        # Pattern 1: Number with decimals (with or without apostrophe) + CHF
        # Matches: "1150.00 CHF", "1'150.00 CHF", "1'150.00 CHF" (U+2019 curly quote)
        if len(chf_int) == 4:  # 4-digit number (could have apostrophe after first digit)
            pattern1 = chf_int[0] + r"['\u2019]?" + chf_int[1:] + r"\.00\s+[Cc][Hh][Ff]"
        elif len(chf_int) == 5:  # 5-digit number (could have apostrophe after 2nd digit)
            pattern1 = chf_int[:2] + r"['\u2019]?" + chf_int[2:] + r"\.00\s+[Cc][Hh][Ff]"
        else:
            pattern1 = re.escape(chf_int) + r"\.00\s+[Cc][Hh][Ff]"

        for match in re.finditer(pattern1, text):
            all_replacements.append((match.start(), match.end(), f'{eur_value} EUR'))

        # Pattern 2: CHF + Number (with or without decimals, with or without apostrophe)
        # Matches: "CHF 1150", "CHF 1150.00", "CHF 1'150", "CHF 1'150.00" (U+2019 curly quote)
        if len(chf_int) == 4:
            pattern2 = r"[Cc][Hh][Ff]\s+" + chf_int[0] + r"['\u2019]?" + chf_int[1:] + r"(?:\.00)?(?!\d)"
        elif len(chf_int) == 5:
            pattern2 = r"[Cc][Hh][Ff]\s+" + chf_int[:2] + r"['\u2019]?" + chf_int[2:] + r"(?:\.00)?(?!\d)"
        else:
            pattern2 = r"[Cc][Hh][Ff]\s+" + re.escape(chf_int) + r"(?:\.00)?(?!\d)"

        for match in re.finditer(pattern2, text):
            all_replacements.append((match.start(), match.end(), f'EUR {eur_value}'))

        # Pattern 3: Number + CHF (with or without decimals, with or without apostrophe)
        # Matches: "1150 CHF", "1150.00 CHF", "1'150 CHF", "1'150.00 CHF" (U+2019 curly quote), "190 CHF"
        if len(chf_int) == 4:
            pattern3 = chf_int[0] + r"['\u2019]?" + chf_int[1:] + r"(?:\.00)?\s+[Cc][Hh][Ff]"
        elif len(chf_int) == 5:
            pattern3 = chf_int[:2] + r"['\u2019]?" + chf_int[2:] + r"(?:\.00)?\s+[Cc][Hh][Ff]"
        else:
            pattern3 = re.escape(chf_int) + r"(?:\.00)?\s+[Cc][Hh][Ff]"

        for match in re.finditer(pattern3, text):
            all_replacements.append((match.start(), match.end(), f'{eur_value} EUR'))

        # Pattern 4: Number.XX directly followed by CHF (no space) - e.g., "33.00CHF"
        # Matches: "33.00CHF", "1150.00CHF", "1'150.00CHF"
        if len(chf_int) == 4:
            pattern4 = chf_int[0] + r"['\u2019]?" + chf_int[1:] + r"\.00[Cc][Hh][Ff]"
        elif len(chf_int) == 5:
            pattern4 = chf_int[:2] + r"['\u2019]?" + chf_int[2:] + r"\.00[Cc][Hh][Ff]"
        else:
            pattern4 = re.escape(chf_int) + r"\.00[Cc][Hh][Ff]"

        for match in re.finditer(pattern4, text):
            all_replacements.append((match.start(), match.end(), f'{eur_value} EUR'))

        # Pattern 5: Magnum followed by number (no CHF indicator) - e.g., "Magnum 52.00"
        # Matches: "Magnum 52.00", "Magnum 1'150.00"
        if len(chf_int) == 4:
            pattern5 = r"\bMagnum\s+" + chf_int[0] + r"['\u2019]?" + chf_int[1:] + r"\.00"
        elif len(chf_int) == 5:
            pattern5 = r"\bMagnum\s+" + chf_int[:2] + r"['\u2019]?" + chf_int[2:] + r"\.00"
        else:
            pattern5 = r"\bMagnum\s+" + re.escape(chf_int) + r"\.00"

        for match in re.finditer(pattern5, text, re.IGNORECASE):
            # Replace "Magnum XX.XX" with "Magnum YY.YY EUR"
            all_replacements.append((match.start(), match.end(), f'Magnum {eur_value} EUR'))

        # Pattern 6: 36x followed by number (no CHF after) - e.g., "36x 99.00 + VAT"
        # Matches: "36x 99.00", "36x 1'150.00" but NOT "36x 33.00CHF"
        if len(chf_int) == 4:
            pattern6 = r"\b36\s*x\s+" + chf_int[0] + r"['\u2019]?" + chf_int[1:] + r"\.00(?![Cc][Hh][Ff])"
        elif len(chf_int) == 5:
            pattern6 = r"\b36\s*x\s+" + chf_int[:2] + r"['\u2019]?" + chf_int[2:] + r"\.00(?![Cc][Hh][Ff])"
        else:
            pattern6 = r"\b36\s*x\s+" + re.escape(chf_int) + r"\.00(?![Cc][Hh][Ff])"

        for match in re.finditer(pattern6, text, re.IGNORECASE):
            # Replace "36x XX.XX" with "36x YY.YY EUR"
            all_replacements.append((match.start(), match.end(), f'36x {eur_value} EUR'))

    # Remove overlapping replacements (keep the longest match at each position)
    all_replacements.sort(key=lambda x: (x[0], -(x[1] - x[0])))  # Sort by start, then by length descending

    unique_replacements = []
    used_positions = set()

    for start, end, replacement in all_replacements:
        # Check if any position in this range is already used
        if not any(pos in used_positions for pos in range(start, end)):
            unique_replacements.append((start, end, replacement))
            # Mark all positions in this range as used
            for pos in range(start, end):
                used_positions.add(pos)

    # Sort by start position (descending) to replace from end to beginning
    unique_replacements.sort(key=lambda x: x[0], reverse=True)

    # Apply replacements from end to start (so positions don't shift)
    for start, end, replacement in unique_replacements:
        new_text = new_text[:start] + replacement + new_text[end:]

    # Replace any remaining standalone "CHF" with "EUR"
    new_text = re.sub(CHF_PATTERN, 'EUR', new_text)

    # Delete existing runs
    for i in range(len(paragraph.runs) - 1, -1, -1):
        p = paragraph._element
        p.remove(paragraph.runs[i]._element)

    # Rebuild the paragraph content
    current_position = 0
    eur_to_info_map = {v[0]: (k, v[1], v[2], v[3]) for k, v in replacements_to_do.items()}

    # Find all occurrences of the new EUR values in the new_text
    # Sort EUR values by length (descending) to find longer matches first and avoid substring issues
    sorted_eur_values = sorted(eur_to_info_map.keys(), key=len, reverse=True)

    positions = []
    used_positions_set = set()

    for eur_value in sorted_eur_values:
        start = 0
        while True:
            start = new_text.find(eur_value, start)
            if start == -1:
                break

            # Check if this position overlaps with any already-used position
            overlaps = False
            for pos in range(start, start + len(eur_value)):
                if pos in used_positions_set:
                    overlaps = True
                    break

            if not overlaps:
                positions.append((start, eur_value))
                # Mark all positions in this range as used
                for pos in range(start, start + len(eur_value)):
                    used_positions_set.add(pos)

            start += len(eur_value)

    positions.sort(key=lambda x: x[0])

    for start_index, eur_value in positions:
        # Add text segment BEFORE the replaced value
        text_before = new_text[current_position:start_index]
        if text_before:
            paragraph.add_run(text_before)

        # Add the HIGHLIGHTED/non-highlighted new value
        _, highlight_color, _, _ = eur_to_info_map[eur_value]

        run = paragraph.add_run(eur_value)

        if highlight_color is not None:
            run.font.highlight_color = highlight_color

        local_replacements += 1
        current_position = start_index + len(eur_value)

    # Add any remaining text after the last replacement
    if current_position < len(new_text):
        paragraph.add_run(new_text[current_position:])

    return local_replacements


def clean_apostrophes_in_numbers(doc):
    """
    Clean and normalize numbers in the document:
    1. Remove apostrophes (both ' and ' and ') from numbers
    2. Fix malformed numbers like 1480.0.00 (double dots) -> 1480.00

    This preprocesses the document to avoid issues with Swiss number formatting.
    Example: 1'500.00 CHF -> 1500.00 CHF
    Example: 1'480'000.00 CHF -> 1480000.00 CHF
    Example: 1150.0.00 EUR -> 1150.00 EUR
    """
    apostrophes_removed = 0
    malformed_fixed = 0

    def clean_number_text(text):
        """Clean a single text string"""
        nonlocal apostrophes_removed, malformed_fixed

        original = text

        # Step 1: Remove ALL apostrophes from numbers (handles multiple apostrophes)
        # Pattern: digit + apostrophe(s) + digit
        # This handles: ', ', ', `, ʼ, etc.
        while re.search(r"(\d)['\u2019\u0027\u02BC\u2018](\d)", text):
            text = re.sub(r"(\d)['\u2019\u0027\u02BC\u2018](\d)", r'\1\2', text)

        # Step 2: Fix malformed decimals like "1150.0.00" or "1480.0.00"
        # Pattern: digit + ".0." + digits (should be just ".digits")
        text = re.sub(r'(\d)\.0\.(\d{2})\b', r'\1.\2', text)

        # Step 3: Fix triple dots or more: "1150...00" -> "1150.00"
        text = re.sub(r'(\d)\.{2,}(\d{2})\b', r'\1.\2', text)

        if text != original:
            if re.search(r"['\u2019\u0027\u02BC\u2018]", original):
                apostrophes_removed += 1
            if re.search(r'\.0\.\d{2}|\.{2,}', original):
                malformed_fixed += 1

        return text

    # Process all paragraphs
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.text = clean_number_text(run.text)

    # Process all tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.text = clean_number_text(run.text)

    if apostrophes_removed > 0:
        print(f"✅ Preprocessed: removed apostrophes from {apostrophes_removed} number(s)")
    if malformed_fixed > 0:
        print(f"✅ Preprocessed: fixed {malformed_fixed} malformed number(s) (e.g., 1150.0.00 → 1150.00)")

    return doc


def main():
    """Main function to orchestrate the conversion process."""

    print("\n" + "="*80)
    print("CHF to EUR Converter with Wine Name Matching")
    print("="*80 + "\n")

    doc, conversion_map, wine_data_map, duplicate_chf_prices, item_number_map = load_data_and_document()

    if not doc:
        print("\nOperation aborted due to file loading errors.")
        return

    # PREPROCESSING: Remove apostrophes from numbers to avoid formatting issues
    doc = clean_apostrophes_in_numbers(doc)

    total_replacements = 0
    all_numbers_found = []
    conversion_stats = {
        'direct': 0,
        'exact_matched': 0,
        'fuzzy_matched': 0,
        'ambiguous': 0,
        'fallback': 0
    }

    # List to track all conversions for Excel export
    conversion_records = []

    # 1. Iterate through all paragraphs in the document
    for paragraph in doc.paragraphs:
        total_replacements += replace_and_highlight(
            paragraph, conversion_map, wine_data_map,
            duplicate_chf_prices, all_numbers_found, conversion_stats, conversion_records, item_number_map
        )

    # 2. Iterate through all tables in the document (if any)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    total_replacements += replace_and_highlight(
                        paragraph, conversion_map, wine_data_map,
                        duplicate_chf_prices, all_numbers_found, conversion_stats, conversion_records, item_number_map
                    )

    # --- Statistics Report ---
    print("\n" + "="*80)
    print("CONVERSION STATISTICS")
    print("="*80)
    print(f"✅ Direct conversions (unique CHF->EUR): {conversion_stats['direct']}")
    print(f"✅ Exact matches (single option): {conversion_stats['exact_matched']}")
    print(f"🟢 Fuzzy matched (wine name + filters): {conversion_stats['fuzzy_matched']} [GREEN highlight]")
    print(f"🟡 Ambiguous (duplicate, poor match): {conversion_stats['ambiguous']} [YELLOW highlight]")
    print(f"🔴 Market prices & Not in Excel (1.08): {conversion_stats['fallback']} [RED highlight]")
    print(f"\nTotal prices processed: {total_replacements}")
    print("\nFilter criteria applied for duplicates (in order):")
    print("  1. Campaign Sub-Type = 'Normal'")
    print("  2. Size = 75 cl")
    print("  3. Minimum Quantity = 0 or 36 (based on '36x' or '36 bottles' context)")
    print("  4. Campaign Type = 'PRIVATE' (only for Min Qty = 0)")
    print("  5. Wine name fuzzy matching (if multiple options remain)")
    print("="*80)

    # --- Duplicate Check Report ---
    duplicate_counts = Counter(all_numbers_found)
    duplicates_found_in_doc = {num: count for num, count in duplicate_counts.items() if count > 1}

    print("\n" + "="*80)
    print("DUPLICATE PRICES IN DOCUMENT")
    print("="*80)
    if duplicates_found_in_doc:
        print("⚠️  Duplicate CHF prices found in document:")
        for num, count in sorted(duplicates_found_in_doc.items(), key=lambda x: float(x[0])):
            status = ""
            if num in duplicate_chf_prices:
                status = " [Has multiple EUR options in Excel]"
            print(f"  - {num} CHF appears {count} times{status}")
    else:
        print("✅ No duplicate prices found in the document.")
    print("="*80)

    # 3. Save the new document
    try:
        doc.save(NEW_WORD_FILE_PATH)
        print(f"\n🎉 SUCCESS! New Word file saved as:\n   {NEW_WORD_FILE_PATH}")
        print("\nHighlight Legend:")
        print("  🟢 GREEN = Matched using wine name proximity")
        print("  🟡 YELLOW = Ambiguous duplicate (used 1.08 formula)")
        print("  🔴 RED = Not in Excel (used 1.08 formula)")
        print("  No highlight = Direct match from Excel")
    except Exception as e:
        print(f"\n❌ Error saving the new Word file: {e}")

    # 4. Export conversion data to Lines.xlsx
    try:
        export_to_lines_excel(conversion_records, LINES_EXCEL_PATH)
    except PermissionError:
        print(f"\n⚠️  Could not save Lines.xlsx - file may be open in Excel.")
        print(f"   Please close {LINES_EXCEL_PATH} and run the converter again.")
    except Exception as e:
        print(f"\n❌ Error exporting to Lines.xlsx: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()

# ==============================================================================
