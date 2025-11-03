# ==============================================================================
# SCRIPT NAME: word_converter_improved.py
# DESCRIPTION: Converts CHF numbers to EUR in a Word document based on an Excel map.
#              Uses fuzzy wine name matching to resolve duplicate CHF prices.
#              Applies 1.08 fallback with integer floor rounding if a number is
#              not found. Highlights converted numbers based on conversion type:
#              - RED: 1.08 fallback conversion
#              - YELLOW: Duplicate CHF with ambiguous EUR conversion (not in Excel)
#              - GREEN: Successfully matched using wine name proximity
# ==============================================================================

import pandas as pd
from docx import Document
import re
from docx.enum.text import WD_COLOR_INDEX
import math
from collections import Counter, defaultdict
from difflib import SequenceMatcher
import sys
import io

# Fix encoding for Windows console
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

# --- CONFIGURATION (UPDATE THESE PATHS) ---
WORD_FILE_PATH = r"C:\Users\Marco.Africani\Desktop\Month recap\month recap.docx"
EXCEL_FILE_PATH = r"C:\Users\Marco.Africani\Desktop\Month recap\Conversion_month.xlsx"
NEW_WORD_FILE_PATH = r"C:\Users\Marco.Africani\Desktop\Month recap\month recap_EUR.docx"

# Excel column names
CHF_COL = 'Unit Price'
EUR_COL = 'Unit Price (EUR)'
WINE_NAME_COL = 'Wine Name'
CAMPAIGN_SUBTYPE_COL = 'Campaign Sub-Type'
CAMPAIGN_TYPE_COL = 'Campaign Type'
SIZE_COL = 'Size'
MIN_QUANTITY_COL = 'Minimum Quantity'
COMPETITOR_CODE_COL = 'Competitor Code'
PRODUCER_NAME_COL = 'Producer Name'
VINTAGE_COL = 'Vintage'

# Define Regex Patterns
# Match numbers with decimals, including Swiss format with apostrophe (e.g., 1'500.00)
NUMBER_PATTERN = r"\d+(?:'\d{3})*\.\d{2}"
CHF_PATTERN = r'[Cc][Hh][Ff]'

# Additional pattern to catch numbers like "CHF 100" or "CHF 42+vat" (without decimals)
CHF_NUMBER_NO_DECIMAL = r"[Cc][Hh][Ff]\s+(\d+(?:'\d{3})*)(?![.\d])"

# Pattern to catch "NUMBER CHF" format without decimals (e.g., "190 CHF")
NUMBER_THEN_CHF = r"(\d+(?:'\d{3})*)\s+[Cc][Hh][Ff]"

# Fuzzy matching threshold (0-1, where 1 is exact match)
FUZZY_MATCH_THRESHOLD = 0.5


def normalize_wine_name(name):
    """
    Normalize wine name for better matching:
    - Convert to lowercase
    - Remove extra whitespace
    - Remove common punctuation
    """
    if not isinstance(name, str):
        return ""

    # Convert to lowercase
    name = name.lower()

    # Remove château/chateau variations and common prefixes
    name = re.sub(r'\bch[âa]teau\b', '', name)
    name = re.sub(r'\bdomaine\b', '', name)
    name = re.sub(r'\bch[âa]teau\b', '', name)

    # Remove special characters but keep spaces
    name = re.sub(r'[^\w\s]', ' ', name)

    # Remove extra whitespace
    name = ' '.join(name.split())

    return name.strip()


def calculate_similarity(text1, text2):
    """
    Calculate similarity ratio between two strings (0-1).
    Uses both full string matching and partial matching for better results.
    """
    text1_norm = normalize_wine_name(text1)
    text2_norm = normalize_wine_name(text2)

    if not text1_norm or not text2_norm:
        return 0.0

    # Full string similarity
    full_similarity = SequenceMatcher(None, text1_norm, text2_norm).ratio()

    # Check if one is contained in the other (partial match)
    if text1_norm in text2_norm or text2_norm in text1_norm:
        # Boost score for substring matches
        full_similarity = max(full_similarity, 0.7)

    # Word-level matching (check if key words are shared)
    words1 = set(text1_norm.split())
    words2 = set(text2_norm.split())

    # Remove common filler words
    filler_words = {'the', 'de', 'di', 'du', 'della', 'des', 'le', 'la', 'del'}
    words1 = words1 - filler_words
    words2 = words2 - filler_words

    if words1 and words2:
        # Jaccard similarity for word sets
        word_overlap = len(words1 & words2) / len(words1 | words2)
        # Combine full string and word-level similarity
        combined_similarity = max(full_similarity, word_overlap * 0.9)
    else:
        combined_similarity = full_similarity

    return combined_similarity


def detect_market_price_context(text, price_match_start):
    """
    Detect if a price is a market price reference (should use 1.08 conversion).
    Returns True if market price indicators are found VERY CLOSE to this specific price.
    """
    # Look in TIGHT context (only 40 chars before, 30 after) to avoid false positives
    context_before = text[max(0, price_match_start - 40):price_match_start].lower()
    context_after = text[price_match_start:min(len(text), price_match_start + 30)].lower()

    # Market price indicators that must be IMMEDIATELY before "CHF"
    # Pattern: "market price stands at approximately CHF 100"
    #          "swiss market price ... CHF 100"
    immediate_before_patterns = [
        r'market\s+price\s+chf\s*$',           # "market price CHF" directly before number
        r'approximately\s+chf\s*$',             # "approximately CHF" directly before
        r'around\s+chf\s*$',                    # "around CHF" directly before
        r'stands\s+at.*chf\s*$',                # "stands at ... CHF" directly before
    ]

    # Check if "CHF" appears right before the price position in original text
    # This helps identify if this is part of "CHF XX" pattern
    chf_before_price = text[max(0, price_match_start - 5):price_match_start]
    has_chf_prefix = bool(re.search(r'chf\s*$', chf_before_price, re.IGNORECASE))

    # Only apply market price detection if CHF prefix exists
    if has_chf_prefix:
        for pattern in immediate_before_patterns:
            if re.search(pattern, context_before):
                return True

    # Market price indicators AFTER the price (in parentheses)
    # Pattern: "100.00 EUR + VAT (market price EUR 108.00)"
    market_indicators_after = [
        r'^\s*\)?\s*\(?\s*market\s+price',     # "(market price" right after
    ]

    for indicator in market_indicators_after:
        if re.search(indicator, context_after):
            return True

    return False


def detect_quantity_indicator(text, price_match_start):
    """
    Detect quantity indicators near the price (e.g., "36x", "36 bottles").
    Returns the minimum quantity (0 or 36) based on context.
    IMPORTANT: Must be VERY close to the price to avoid false positives.
    """
    # Look in TIGHT context around the price
    context_before = text[max(0, price_match_start - 30):price_match_start].lower()
    context_after = text[price_match_start:min(len(text), price_match_start + 35)].lower()

    # Check for 36-bottle indicators (MUST BE VERY CLOSE)
    # Pattern 1: "36x" directly before the price (within 10 chars)
    if re.search(r'36\s*x\b', context_before[-10:]):
        return 36

    # Pattern 2: "36+ bottle" directly before or after (within 25 chars)
    if re.search(r'36\s*\+\s*bottle', context_before):
        return 36
    if re.search(r'^[^a-z]*36\s*\+\s*bottle', context_after):
        return 36

    # Pattern 3: "if you take 36 bottles" AFTER the price (not before!)
    # This must be in the context_after, not before
    if re.search(r'if\s+you\s+take\s+36\s+bottle', context_after):
        return 36

    # Pattern 4: Price at end of "36 bottles" phrase (within 20 chars before)
    if re.search(r'36\s+bottles?\s*(at|for)?\s*chf', context_before):
        return 36

    # Default to 0 (normal price, no minimum quantity)
    return 0


def extract_vintage_from_context(text, price_match_start):
    """
    Extract vintage year from context around the price.
    Returns integer year or None.
    """
    # Look in wider context
    context = text[max(0, price_match_start - 300):min(len(text), price_match_start + 100)]

    # Find 4-digit years (vintage years typically 1990-2030)
    year_matches = re.findall(r'\b(19[9]\d|20[0-3]\d)\b', context)

    if year_matches:
        # Return the most recent/last mentioned year
        try:
            return int(year_matches[-1])
        except:
            return None

    return None


def extract_wine_name_from_context(text, price_match_start):
    """
    Extract potential wine name from context around the price.
    Looks backwards from the price position to find wine name.
    """
    # Get text before the price (up to 400 characters for better context)
    context_before = text[max(0, price_match_start - 400):price_match_start]

    wine_candidates = []

    # Pattern 1: FIRST colon in the context (usually the wine name at paragraph start)
    # e.g., "Château Rieussec 2019: ... price"
    # Find ALL colons, prefer the FIRST one (which is usually the wine name)
    all_colons = list(re.finditer(r'([A-ZÀ-ÿ][^\n:]{3,60})[:]\s*', context_before))
    if all_colons:
        # Take the FIRST colon match (likely the wine name)
        first_colon = all_colons[0]
        candidate = first_colon.group(1).strip()
        # Remove common trailing words
        candidate = re.sub(r'\s+(at|from|for|with|the|a|an)$', '', candidate, flags=re.IGNORECASE)
        wine_candidates.append(candidate)

        # Also consider the LAST colon if different (might be more specific context)
        if len(all_colons) > 1:
            last_colon = all_colons[-1]
            candidate_last = last_colon.group(1).strip()
            candidate_last = re.sub(r'\s+(at|from|for|with|the|a|an)$', '', candidate_last, flags=re.IGNORECASE)
            # Only add if it's different and doesn't contain "CHF" or price-related words
            if candidate_last != candidate and not re.search(r'chf|price|offer', candidate_last, re.IGNORECASE):
                wine_candidates.append(candidate_last)

    # Pattern 2: Quoted text (wine names often in quotes)
    # e.g., "the famous "Château Pavie" at"
    quote_matches = re.findall(r'["""]([^"""]{3,60})["""]', context_before)
    for match in quote_matches:
        # Prefer quotes with capitalized content
        if re.search(r'[A-ZÀ-ÿ]', match):
            wine_candidates.append(match.strip())

    # Pattern 3: Château/Domaine followed by name
    # e.g., "Château Montrose 2021"
    chateau_pattern = re.findall(r'\b([CcDd]h[âa]teau|Domaine|Dom\.)\s+([A-ZÀ-ÿ][^\n:,.]{3,40})', context_before)
    for prefix, name in chateau_pattern:
        wine_candidates.append(f"{prefix} {name}".strip())

    # Pattern 4: Producer name patterns (e.g., "Penfolds 2019", "Aalto 2023")
    # Match: Capitalized word(s) optionally followed by year
    producer_pattern = re.findall(
        r'\b([A-ZÀ-ÿ][a-zà-ÿ]+(?:\s+[A-ZÀ-ÿ][a-zà-ÿ]+){0,3})\s+(?:\d{4})?',
        context_before
    )
    if producer_pattern:
        # Get last few capitalized phrases
        wine_candidates.extend(producer_pattern[-3:])

    # Pattern 5: Text between line start and dash/colon
    # e.g., "Aalto 2023: ..." or "Dominus 2016: ..."
    line_start = context_before.split('\n')[-1] if '\n' in context_before else context_before
    line_pattern = re.match(r'^([A-ZÀ-ÿ][^\n:–-]{3,60}?)[:–-]', line_start.strip())
    if line_pattern:
        wine_candidates.append(line_pattern.group(1).strip())

    # Filter and clean candidates
    cleaned_candidates = []
    for candidate in wine_candidates:
        # Remove year patterns at the end
        candidate = re.sub(r'\s+\d{4}\s*$', '', candidate)
        # Remove common noise words at the end
        candidate = re.sub(r'\s+(at|from|for|with|price|the|a|an)$', '', candidate, flags=re.IGNORECASE)
        # Remove extra whitespace
        candidate = ' '.join(candidate.split())
        if len(candidate) >= 3:
            cleaned_candidates.append(candidate)

    # Return the longest and most specific candidate
    if cleaned_candidates:
        # Prefer longer candidates that likely contain more specific info
        return max(cleaned_candidates, key=lambda x: (len(x), x.count(' ')))

    return ""


def load_data_and_document():
    """Loads the Excel data, creates the conversion map, and loads the Word document."""
    conversion_map = {}
    wine_data_map = defaultdict(list)  # CHF -> list of full row data
    duplicate_chf_prices = set()
    doc = None
    df_full = None  # Store full dataframe for reference

    # 1. Load Excel File and Create Conversion Dictionary
    try:
        df = pd.read_excel(EXCEL_FILE_PATH)
        df_full = df.copy()

        # Standardize columns
        df['CHF_KEY_FORMATTED'] = df[CHF_COL].astype(float).round(2).apply(lambda x: f'{x:.2f}')
        df['EUR_VALUE_FORMATTED'] = df[EUR_COL].astype(float).round(2).apply(lambda x: f'{x:.2f}')
        df['WINE_NAME_NORMALIZED'] = df[WINE_NAME_COL].astype(str)

        # Find duplicate CHF prices (same CHF, different EUR)
        chf_eur_mapping = df.groupby('CHF_KEY_FORMATTED')['EUR_VALUE_FORMATTED'].apply(set).to_dict()
        for chf, eur_set in chf_eur_mapping.items():
            if len(eur_set) > 1:
                duplicate_chf_prices.add(chf)

        # Create wine data mapping with all relevant columns
        for _, row in df.iterrows():
            chf = row['CHF_KEY_FORMATTED']
            eur = row['EUR_VALUE_FORMATTED']
            wine = row['WINE_NAME_NORMALIZED']

            # Extract additional columns for filtering
            campaign_subtype = str(row.get(CAMPAIGN_SUBTYPE_COL, '')).strip().lower()
            campaign_type = str(row.get(CAMPAIGN_TYPE_COL, '')).strip().lower()
            size = row.get(SIZE_COL, 0)
            min_qty = row.get(MIN_QUANTITY_COL, 0)
            competitor_code = row.get(COMPETITOR_CODE_COL, None)
            producer_name = str(row.get(PRODUCER_NAME_COL, '')).strip()
            vintage = row.get(VINTAGE_COL, None)

            wine_data_map[chf].append({
                'wine_name': wine,
                'eur_value': eur,
                'campaign_subtype': campaign_subtype,
                'campaign_type': campaign_type,
                'size': size,
                'min_quantity': min_qty,
                'competitor_code': competitor_code,
                'producer_name': producer_name,
                'vintage': vintage
            })

        # For non-duplicate prices, create simple conversion map
        for chf, eur_set in chf_eur_mapping.items():
            if len(eur_set) == 1:
                conversion_map[chf] = list(eur_set)[0]

        print("✅ Excel file loaded successfully.")
        print(f"   - Found {len(conversion_map)} unique CHF->EUR conversions")
        print(f"   - Found {len(duplicate_chf_prices)} duplicate CHF prices requiring wine name matching")
        print(f"   - Loaded {len(df)} rows with full metadata (Campaign Type, Size, Min Quantity)")

    except FileNotFoundError:
        print(f"❌ Error: Excel file not found at {EXCEL_FILE_PATH}")
    except KeyError as e:
        print(f"❌ Error: Column not found. Details: {e}")
    except Exception as e:
        print(f"❌ Unexpected error loading Excel file: {e}")

    # 2. Load Word Document
    try:
        doc = Document(WORD_FILE_PATH)
        print(f"✅ Word file loaded from {WORD_FILE_PATH}")
    except FileNotFoundError:
        print(f"❌ Error: Word file not found at {WORD_FILE_PATH}")
    except Exception as e:
        print(f"❌ Unexpected error loading Word file: {e}")

    return doc, conversion_map, wine_data_map, duplicate_chf_prices


def round_to_5_or_0(value):
    """
    Round value up to the nearest number ending in 5 or 0.
    Used for prices above 300 CHF.
    Examples: 1162 -> 1165, 1163 -> 1165, 1167 -> 1170
    """
    # Get the last digit
    last_digit = int(value) % 10

    if last_digit == 0 or last_digit == 5:
        return value  # Already ends in 0 or 5
    elif last_digit < 5:
        # Round up to 5 (e.g., 1162 -> 1165)
        return int(value) + (5 - last_digit)
    else:
        # Round up to next 10 (e.g., 1167 -> 1170)
        return int(value) + (10 - last_digit)


def find_best_wine_match(chf_price, context_wine_name, wine_data_map, detected_quantity=0,
                         context_vintage=None, context_producer=None):
    """
    Find the best EUR conversion for a CHF price using wine name matching and filtering.

    Args:
        chf_price: CHF price string (e.g., "42.00")
        context_wine_name: Wine name extracted from context
        wine_data_map: Dictionary mapping CHF to list of wine data dicts
        detected_quantity: Detected minimum quantity (0 or 36)
        context_vintage: Vintage year extracted from context (optional)
        context_producer: Producer name extracted from context (optional)

    Returns (eur_value, match_quality)
    """
    if chf_price not in wine_data_map:
        return None, 'not_found'

    wine_options = wine_data_map[chf_price]

    # If only one option, use it
    if len(wine_options) == 1:
        return wine_options[0]['eur_value'], 'exact'

    # Multiple options - apply filters first
    # Filter 1: Campaign Sub-Type = "Normal"
    filtered_options = [opt for opt in wine_options if opt['campaign_subtype'] == 'normal']

    if not filtered_options:
        filtered_options = wine_options

    # Filter 2: Size = 75 (for 75cl bottles)
    size_75_options = [opt for opt in filtered_options if opt['size'] == 75.0 or opt['size'] == 75]

    if size_75_options:
        filtered_options = size_75_options

    # Filter 3: Match quantity (0 for normal, 36 for bulk)
    quantity_matched_options = [opt for opt in filtered_options if opt['min_quantity'] == detected_quantity]

    if quantity_matched_options:
        filtered_options = quantity_matched_options

        # VALIDATION: If detected_quantity == 36, ensure the price is different from normal (qty=0)
        # This prevents matching a normal price when "36 bottles" is mentioned
        if detected_quantity == 36:
            # Get all normal prices (qty=0) for this CHF value
            normal_prices = [opt for opt in wine_options
                           if opt['min_quantity'] == 0 and opt['size'] == 75.0]
            if normal_prices:
                # Get EUR values for normal prices
                normal_eur_values = set(opt['eur_value'] for opt in normal_prices)
                # Filter out any 36-bottle options that have the same EUR as normal
                validated_options = [opt for opt in filtered_options
                                   if opt['eur_value'] not in normal_eur_values]
                if validated_options:
                    filtered_options = validated_options

    # Filter 4: Competitor Code is empty (NaN) - PREFER these rows
    empty_competitor_options = [opt for opt in filtered_options
                                if pd.isna(opt['competitor_code']) or opt['competitor_code'] == '']

    if empty_competitor_options:
        filtered_options = empty_competitor_options

    # Filter 5: Campaign Type = "PRIVATE" (only for quantity=0)
    if len(filtered_options) > 1 and detected_quantity == 0:
        private_options = [opt for opt in filtered_options if opt['campaign_type'] == 'private']
        if private_options:
            filtered_options = private_options

    # Filter 6: Match vintage if available
    if context_vintage and len(filtered_options) > 1:
        vintage_matched = [opt for opt in filtered_options
                          if opt['vintage'] == context_vintage or opt['vintage'] == float(context_vintage)]
        if vintage_matched:
            filtered_options = vintage_matched

    # NEW: Filter 7: Prefer LAST row (bottom-most) when multiple Size=75 matches remain
    # This helps when there are multiple identical wines with different campaigns
    if len(filtered_options) > 1:
        # Add row index to each option (stored during data loading)
        # For now, we'll use the last item in the list as it represents the bottom-most row
        # since pandas iterates top-to-bottom and we append in order
        filtered_options = [filtered_options[-1]]

    # After filtering, if only one option remains, use it
    if len(filtered_options) == 1:
        return filtered_options[0]['eur_value'], 'fuzzy_filtered'

    # Multiple options remain - use intelligent matching
    if not context_wine_name and not context_producer:
        # Use price proximity (choose EUR closest to CHF * 1.08)
        if filtered_options:
            chf_float = float(chf_price)
            expected_eur = chf_float * 1.08
            best_option = min(filtered_options, key=lambda opt: abs(float(opt['eur_value']) - expected_eur))
            return best_option['eur_value'], 'price_proximity'
        return None, 'ambiguous'

    # Calculate similarity scores for each option
    scored_options = []
    for option in filtered_options:
        score = 0.0

        # Wine name similarity
        if context_wine_name:
            wine_similarity = calculate_similarity(context_wine_name, option['wine_name'])
            score += wine_similarity * 2.0  # Weight: 2.0

        # Producer name similarity
        if context_producer and option['producer_name']:
            producer_similarity = calculate_similarity(context_producer, option['producer_name'])
            score += producer_similarity * 1.5  # Weight: 1.5

        # Price proximity bonus (prefer EUR closest to CHF * 1.08)
        chf_float = float(chf_price)
        expected_eur = chf_float * 1.08
        eur_float = float(option['eur_value'])
        price_diff = abs(eur_float - expected_eur)
        # Normalize: closer = higher score
        price_proximity_score = max(0, 1.0 - (price_diff / expected_eur))
        score += price_proximity_score * 0.5  # Weight: 0.5

        scored_options.append((option, score))

    # Sort by score
    scored_options.sort(key=lambda x: x[1], reverse=True)
    best_option, best_score = scored_options[0]

    # If good match (score > threshold), use it
    if best_score >= 1.0:  # Threshold for combined score
        return best_option['eur_value'], 'fuzzy'

    # Otherwise, use price proximity
    chf_float = float(chf_price)
    expected_eur = chf_float * 1.08
    best_by_price = min(filtered_options, key=lambda opt: abs(float(opt['eur_value']) - expected_eur))
    return best_by_price['eur_value'], 'price_proximity'


def replace_and_highlight(paragraph, conversion_map, wine_data_map, duplicate_chf_prices,
                         all_numbers_found, conversion_stats):
    """
    Performs search and replace at the run level, applying highlighting based on conversion type.
    """

    local_replacements = 0

    # 1. Replace CHF with EUR in the paragraph text
    original_text = paragraph.text

    text = original_text

    # Find all number matches with their positions (three patterns)
    # Pattern 1: Standard XX.XX format (including Swiss format with apostrophe like 1'500.00)
    matches_with_positions = []
    for m in re.finditer(NUMBER_PATTERN, text):
        # Remove apostrophes for Excel lookup (1'500.00 -> 1500.00)
        number_clean = m.group().replace("'", "")
        matches_with_positions.append((number_clean, m.start()))

    # Pattern 2: "CHF 100" style (convert to "100.00" format)
    chf_no_decimal_matches = re.finditer(CHF_NUMBER_NO_DECIMAL, text)
    for match in chf_no_decimal_matches:
        number = match.group(1).replace("'", "")  # Remove apostrophes
        formatted_number = f"{number}.00"
        # Add to matches with the position of the number (not CHF)
        matches_with_positions.append((formatted_number, match.start() + match.group().find(match.group(1))))

    # Pattern 3: "190 CHF" style (NUMBER THEN CHF, without decimals)
    number_then_chf_matches = re.finditer(NUMBER_THEN_CHF, text)
    for match in number_then_chf_matches:
        number = match.group(1).replace("'", "")  # Remove apostrophes
        # Check if this number already has a decimal match (skip if so)
        if not any(chf == f"{number}.00" or chf == number for chf, _ in matches_with_positions):
            formatted_number = f"{number}.00"
            matches_with_positions.append((formatted_number, match.start()))

    # Dictionary of replacements: {chf_str: (new_eur_str, highlight_color_index, context)}
    replacements_to_do = {}

    for chf_str, position in matches_with_positions:

        # --- VINTAGE/YEAR SKIP CHECK ---
        try:
            float_val = float(chf_str)
            if float_val.is_integer() and 1000 <= int(float_val) <= 9999:
                continue
        except ValueError:
            continue

        all_numbers_found.append(chf_str)

        eur_value = None
        highlight_color = None
        match_type = None

        # Detect market price context (should force 1.08 conversion)
        is_market_price = detect_market_price_context(text, position)

        # Extract context information
        context_wine = extract_wine_name_from_context(text, position)
        context_vintage = extract_vintage_from_context(text, position)

        # Try to extract producer name from wine name
        context_producer = None
        if context_wine:
            # Check if wine name contains producer patterns
            producer_match = re.search(r'(Aalto|Colgin|Roederer|Louis Roederer)', context_wine, re.IGNORECASE)
            if producer_match:
                context_producer = producer_match.group(1)

        # Detect quantity indicator (0 or 36)
        detected_quantity = detect_quantity_indicator(text, position)

        # If it's a market price reference, always use 1.08 conversion with RED highlight
        if is_market_price:
            try:
                raw_eur = float(chf_str) * 1.08
                floored_eur = math.floor(raw_eur)
                eur_value = f"{floored_eur:.2f}"
                highlight_color = WD_COLOR_INDEX.RED
                match_type = 'market_price_1.08'
                conversion_stats['fallback'] += 1
            except ValueError:
                continue

        # Check if it's a known non-duplicate price
        elif chf_str in conversion_map:
            eur_value = conversion_map[chf_str]

            # Apply rounding rule for prices above 300 CHF
            chf_float = float(chf_str)
            if chf_float > 300:
                eur_float = float(eur_value)
                rounded_eur = round_to_5_or_0(eur_float)
                eur_value = f"{rounded_eur:.2f}"

            match_type = 'direct'
            conversion_stats['direct'] += 1

        # Check if it's a duplicate requiring wine name matching
        elif chf_str in duplicate_chf_prices:
            eur_value, quality = find_best_wine_match(
                chf_str, context_wine, wine_data_map, detected_quantity,
                context_vintage, context_producer
            )

            if quality == 'fuzzy' or quality == 'fuzzy_filtered' or quality == 'price_proximity':
                # Apply rounding rule for prices above 300 CHF
                chf_float = float(chf_str)
                if chf_float > 300 and eur_value:
                    eur_float = float(eur_value)
                    rounded_eur = round_to_5_or_0(eur_float)
                    eur_value = f"{rounded_eur:.2f}"

                highlight_color = WD_COLOR_INDEX.GREEN
                match_type = quality
                conversion_stats['fuzzy_matched'] += 1
            elif quality == 'exact':
                # Apply rounding rule for prices above 300 CHF
                chf_float = float(chf_str)
                if chf_float > 300 and eur_value:
                    eur_float = float(eur_value)
                    rounded_eur = round_to_5_or_0(eur_float)
                    eur_value = f"{rounded_eur:.2f}"

                match_type = 'exact_match'
                conversion_stats['exact_matched'] += 1
            elif quality == 'ambiguous':
                # Apply 1.08 fallback and highlight yellow
                try:
                    raw_eur = float(chf_str) * 1.08
                    floored_eur = math.floor(raw_eur)

                    # Apply rounding rule for prices above 300 CHF
                    chf_float = float(chf_str)
                    if chf_float > 300:
                        floored_eur = round_to_5_or_0(floored_eur)

                    eur_value = f"{floored_eur:.2f}"
                    highlight_color = WD_COLOR_INDEX.YELLOW
                    match_type = 'ambiguous'
                    conversion_stats['ambiguous'] += 1
                except ValueError:
                    continue

        # Not in Excel at all - use 1.08 fallback
        else:
            try:
                raw_eur = float(chf_str) * 1.08
                floored_eur = math.floor(raw_eur)

                # Apply rounding rule for prices above 300 CHF
                chf_float = float(chf_str)
                if chf_float > 300:
                    floored_eur = round_to_5_or_0(floored_eur)

                eur_value = f"{floored_eur:.2f}"
                highlight_color = WD_COLOR_INDEX.RED
                match_type = 'fallback_1.08'
                conversion_stats['fallback'] += 1
            except ValueError:
                continue

        if eur_value:
            replacements_to_do[chf_str] = (eur_value, highlight_color, context_wine, match_type)

    # 2. Perform number replacement with highlighting (rebuilding runs)
    new_text = text

    # Build a comprehensive replacement map that handles all formats
    all_replacements = []  # List of (original_text, new_text, position)

    for chf_str, (eur_value, _, _, _) in replacements_to_do.items():
        chf_number_only = chf_str.replace('.00', '')

        # Format 1: "CHF 1500" or "CHF 1'500" -> "EUR {eur_value}"
        pattern1 = r'[Cc][Hh][Ff]\s+' + re.escape(chf_number_only).replace('\\', '') + r"(?:'\d{3})*(?!\.\d)"
        for match in re.finditer(pattern1, text):
            all_replacements.append((match.start(), match.end(), f'EUR {eur_value}'))

        # Format 2: "1500 CHF" or "1'500 CHF" -> "{eur_value} EUR"
        pattern2 = re.escape(chf_number_only).replace('\\', '') + r"(?:'\d{3})*\s+[Cc][Hh][Ff]"
        for match in re.finditer(pattern2, text):
            all_replacements.append((match.start(), match.end(), f'{eur_value} EUR'))

        # Format 3: "1500.00" or "1'500.00" -> "{eur_value}"
        # Try with apostrophe first
        chf_with_apostrophe = re.sub(r'(\d)(\d{3})\.', r"\1'\2.", chf_str)
        if chf_with_apostrophe in text:
            for idx in range(len(text)):
                if text[idx:idx+len(chf_with_apostrophe)] == chf_with_apostrophe:
                    all_replacements.append((idx, idx+len(chf_with_apostrophe), eur_value))

        # Try without apostrophe
        if chf_str in text:
            for idx in range(len(text)):
                if text[idx:idx+len(chf_str)] == chf_str:
                    all_replacements.append((idx, idx+len(chf_str), eur_value))

    # Sort by position (start index) in reverse order to replace from end to start
    all_replacements.sort(key=lambda x: x[0], reverse=True)

    # Remove overlapping replacements (keep first = highest priority)
    seen_ranges = []
    unique_replacements = []
    for start, end, replacement in all_replacements:
        # Check if this range overlaps with any seen range
        overlaps = False
        for seen_start, seen_end in seen_ranges:
            if not (end <= seen_start or start >= seen_end):
                overlaps = True
                break
        if not overlaps:
            unique_replacements.append((start, end, replacement))
            seen_ranges.append((start, end))

    # Apply replacements from end to start
    for start, end, replacement in unique_replacements:
        new_text = new_text[:start] + replacement + new_text[end:]

    # Replace any remaining CHF with EUR
    new_text = re.sub(CHF_PATTERN, 'EUR', new_text)

    # Delete existing runs
    for i in range(len(paragraph.runs) - 1, -1, -1):
        p = paragraph._element
        p.remove(paragraph.runs[i]._element)

    # Rebuild the paragraph content
    current_position = 0
    eur_to_info_map = {v[0]: (k, v[1], v[2], v[3]) for k, v in replacements_to_do.items()}

    # Find all occurrences of the new EUR values in the new_text
    positions = []
    for eur_value in eur_to_info_map.keys():
        start = 0
        while True:
            start = new_text.find(eur_value, start)
            if start == -1:
                break
            positions.append((start, eur_value))
            start += len(eur_value)

    positions.sort(key=lambda x: x[0])

    for start_index, eur_value in positions:
        # Add text segment BEFORE the replaced value
        text_before = new_text[current_position:start_index]
        if text_before:
            paragraph.add_run(text_before)

        # Add the HIGHLIGHTED/non-highlighted new value
        _, highlight_color, _, _ = eur_to_info_map[eur_value]

        run = paragraph.add_run(eur_value)

        if highlight_color is not None:
            run.font.highlight_color = highlight_color

        local_replacements += 1
        current_position = start_index + len(eur_value)

    # Add any remaining text after the last replacement
    if current_position < len(new_text):
        paragraph.add_run(new_text[current_position:])

    return local_replacements


def main():
    """Main function to orchestrate the conversion process."""

    print("\n" + "="*80)
    print("CHF to EUR Converter with Wine Name Matching")
    print("="*80 + "\n")

    doc, conversion_map, wine_data_map, duplicate_chf_prices = load_data_and_document()

    if not doc:
        print("\nOperation aborted due to file loading errors.")
        return

    total_replacements = 0
    all_numbers_found = []
    conversion_stats = {
        'direct': 0,
        'exact_matched': 0,
        'fuzzy_matched': 0,
        'ambiguous': 0,
        'fallback': 0
    }

    # 1. Iterate through all paragraphs in the document
    for paragraph in doc.paragraphs:
        total_replacements += replace_and_highlight(
            paragraph, conversion_map, wine_data_map,
            duplicate_chf_prices, all_numbers_found, conversion_stats
        )

    # 2. Iterate through all tables in the document (if any)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    total_replacements += replace_and_highlight(
                        paragraph, conversion_map, wine_data_map,
                        duplicate_chf_prices, all_numbers_found, conversion_stats
                    )

    # --- Statistics Report ---
    print("\n" + "="*80)
    print("CONVERSION STATISTICS")
    print("="*80)
    print(f"✅ Direct conversions (unique CHF->EUR): {conversion_stats['direct']}")
    print(f"✅ Exact matches (single option): {conversion_stats['exact_matched']}")
    print(f"🟢 Fuzzy matched (wine name + filters): {conversion_stats['fuzzy_matched']} [GREEN highlight]")
    print(f"🟡 Ambiguous (duplicate, poor match): {conversion_stats['ambiguous']} [YELLOW highlight]")
    print(f"🔴 Market prices & Not in Excel (1.08): {conversion_stats['fallback']} [RED highlight]")
    print(f"\nTotal prices processed: {total_replacements}")
    print("\nFilter criteria applied for duplicates (in order):")
    print("  1. Campaign Sub-Type = 'Normal'")
    print("  2. Size = 75 cl")
    print("  3. Minimum Quantity = 0 or 36 (based on '36x' or '36 bottles' context)")
    print("  4. Campaign Type = 'PRIVATE' (only for Min Qty = 0)")
    print("  5. Wine name fuzzy matching (if multiple options remain)")
    print("="*80)

    # --- Duplicate Check Report ---
    duplicate_counts = Counter(all_numbers_found)
    duplicates_found_in_doc = {num: count for num, count in duplicate_counts.items() if count > 1}

    print("\n" + "="*80)
    print("DUPLICATE PRICES IN DOCUMENT")
    print("="*80)
    if duplicates_found_in_doc:
        print("⚠️  Duplicate CHF prices found in document:")
        for num, count in sorted(duplicates_found_in_doc.items(), key=lambda x: float(x[0])):
            status = ""
            if num in duplicate_chf_prices:
                status = " [Has multiple EUR options in Excel]"
            print(f"  - {num} CHF appears {count} times{status}")
    else:
        print("✅ No duplicate prices found in the document.")
    print("="*80)

    # 3. Save the new document
    try:
        doc.save(NEW_WORD_FILE_PATH)
        print(f"\n🎉 SUCCESS! New Word file saved as:\n   {NEW_WORD_FILE_PATH}")
        print("\nHighlight Legend:")
        print("  🟢 GREEN = Matched using wine name proximity")
        print("  🟡 YELLOW = Ambiguous duplicate (used 1.08 formula)")
        print("  🔴 RED = Not in Excel (used 1.08 formula)")
        print("  No highlight = Direct match from Excel")
    except Exception as e:
        print(f"\n❌ Error saving the new Word file: {e}")


if __name__ == "__main__":
    main()

# ==============================================================================
